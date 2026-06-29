# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 读取原始文档
print("正在读取原始 DOCX 文档...")
doc = docx.Document(r'E:\郑州录音\新建 DOCX 文档.docx')

# 提取所有段落文本
raw_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        raw_text.append(text)

print(f"共提取 {len(raw_text)} 个段落")

# 清理录音停顿标记
def clean_text(text):
    # 移除 [...Xs] 这样的停顿标记
    text = re.sub(r'$$\.\.\.\d+\.?\d*s$$', '', text)
    # 移除其他可能的标记
    text = re.sub(r'$$\d+%$$', '', text)
    text = re.sub(r'$$\d+$$|$\d+$', '', text)
    # 清理多余空格
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# 清理所有文本
cleaned_paragraphs = []
for text in raw_text:
    cleaned = clean_text(text)
    if cleaned:
        cleaned_paragraphs.append(cleaned)

print(f"清理后共 {len(cleaned_paragraphs)} 个段落")

# 创建新的学习资料文档
new_doc = docx.Document()

# 设置文档标题
title = new_doc.add_heading('外贸全链条培训学习资料', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加说明
new_doc.add_paragraph('本文档基于实战培训录音整理，涵盖外贸全链条运营、中医出海策略、品牌运营官技能等核心内容，适合外贸从业者、跨境电商运营人员系统学习。')
new_doc.add_paragraph()

# 根据内容逻辑分章节
# 分析内容，识别关键点
chapters = {
    '第一章：外贸全链条概述': [],
    '第二章：中医出海市场机遇': [],
    '第三章：品牌运营官的角色定位': [],
    '第四章：客户开发策略（TikTok, Facebook, WhatsApp）': [],
    '第五章：报价与谈判技巧': [],
    '第六章：供应链管理': [],
    '第七章：跨文化沟通与礼仪': [],
    '第八章：实战案例解析': [],
    '第九章：常见问题与解决方案': [],
    '第十章：行动计划与总结': []
}

# 简单的内容分类逻辑（根据关键词）
chapter_keywords = {
    '第一章：外贸全链条概述': ['外贸', '全链条', '概述', '流程', '基本'],
    '第二章：中医出海市场机遇': ['中医', '出海', '市场', '机遇', '趋势', '海外'],
    '第三章：品牌运营官的角色定位': ['品牌', '运营官', '角色', '定位', '职责'],
    '第四章：客户开发策略（TikTok, Facebook, WhatsApp）': ['TikTok', 'Facebook', 'WhatsApp', '客户', '开发', '社媒', '社交'],
    '第五章：报价与谈判技巧': ['报价', '谈判', '价格', '成本', '利润'],
    '第六章：供应链管理': ['供应链', '工厂', '生产', '质量', '交期'],
    '第七章：跨文化沟通与礼仪': ['文化', '沟通', '礼仪', '差异', '跨文'],
    '第八章：实战案例解析': ['案例', '实战', '解析', '分析', '例子'],
    '第九章：常见问题与解决方案': ['问题', '解决', '方案', '常见', '误区'],
    '第十章：行动计划与总结': ['行动', '计划', '总结', '下一步', '建议']
}

# 将段落分配到章节
current_chapter = '第一章：外贸全链条概述'
for para_text in cleaned_paragraphs:
    # 检查是否包含章节关键词
    assigned = False
    for chapter, keywords in chapter_keywords.items():
        if any(keyword in para_text for keyword in keywords):
            chapters[chapter].append(para_text)
            assigned = True
            break
    
    # 如果没有匹配到关键词，放到当前章节
    if not assigned:
        chapters[current_chapter].append(para_text)
    
    # 更新当前章节（简单逻辑：如果段落较短且包含"第X章"字样）
    if len(para_text) < 50 and ('第一章' in para_text or '第二章' in para_text or '第三章' in para_text):
        for ch in chapters.keys():
            if ch in para_text:
                current_chapter = ch
                break

# 生成文档内容
for chapter_title, content_list in chapters.items():
    if content_list:  # 只添加有内容的章节
        # 添加章节标题
        new_doc.add_heading(chapter_title, 1)
        
        # 添加章节内容
        for para_text in content_list:
            new_doc.add_paragraph(para_text)
        
        # 章节间添加空行
        new_doc.add_paragraph()

# 保存文档
output_path = r'E:\郑州录音\外贸全链条培训学习资料_最终版.docx'
new_doc.save(output_path)
print(f"✅ 最终学习资料已保存到: {output_path}")
print(f"📊 统计信息:")
print(f"   - 总章节数: {len([k for k, v in chapters.items() if v])}")
print(f"   - 总段落数: {sum(len(v) for v in chapters.values())}")
