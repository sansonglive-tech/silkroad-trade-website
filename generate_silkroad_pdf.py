#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
新丝路跨境 - 实战出海 - 第二套商业模式方案书
Silk Road themed PDF generation using reportlab
"""

import sys, os, io, math
from collections import defaultdict

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor, white
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (Paragraph, Spacer, Table, TableStyle,
                                 PageBreak, Frame, PageTemplate,
                                 BaseDocTemplate, NextPageTemplate)
from reportlab.platypus.flowables import Flowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas as pdfcanvas

# ═══════════════════════════════════════════════════════
#  Color Palette
# ═══════════════════════════════════════════════════════
GOLD       = HexColor('#C9A96E')
DEEP_BLUE  = HexColor('#1B3A5C')
BRICK_RED  = HexColor('#A0522D')
CREAM      = HexColor('#FFF8E7')
INK_BLACK  = HexColor('#2C2C2C')
LIGHT_GOLD = HexColor('#E8D5A8')
DARK_GOLD  = HexColor('#B8974C')

# ═══════════════════════════════════════════════════════
#  Font Registration
# ═══════════════════════════════════════════════════════
def register_fonts():
    yahei = 'C:/Windows/Fonts/msyh.ttc'
    yahei_bd = 'C:/Windows/Fonts/msyhbd.ttc'
    simsun = 'C:/Windows/Fonts/simsun.ttc'
    if os.path.exists(yahei):
        pdfmetrics.registerFont(TTFont('msyh', yahei))
    elif os.path.exists(simsun):
        pdfmetrics.registerFont(TTFont('msyh', simsun))
    else:
        raise RuntimeError("No Chinese font found")
    if os.path.exists(yahei_bd):
        pdfmetrics.registerFont(TTFont('msyhbd', yahei_bd))
    else:
        pdfmetrics.registerFont(TTFont('msyhbd', yahei))

# ═══════════════════════════════════════════════════════
#  Silk Road Decoration Drawing
# ═══════════════════════════════════════════════════════

def draw_corner_geom(c, x, y, size):
    """Draw a geometric ornament at a corner."""
    c.saveState()
    c.translate(x, y)
    c.setStrokeColor(GOLD)
    c.setLineWidth(1.0)
    half = size / 2
    # Diamond
    p = c.beginPath()
    p.moveTo(0, -half)
    p.lineTo(half, 0)
    p.lineTo(0, half)
    p.lineTo(-half, 0)
    p.lineTo(0, -half)
    c.drawPath(p, stroke=1, fill=0)
    # Inner square
    inner = half * 0.5
    c.rect(-inner, -inner, inner*2, inner*2)
    c.restoreState()

def draw_arabesque_band(c, x, y, w, h=6*mm):
    """Wave/arc decorative band."""
    c.saveState()
    c.translate(x, y)
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.6)
    n = max(4, int(w / (5*mm)))
    sw = w / n
    for i in range(n):
        cx = i * sw + sw / 2
        if i % 2 == 0:
            c.arc(cx - sw/2, -h/2, cx + sw/2, h/2, 0, 180)
        else:
            c.arc(cx - sw/2, -h/2, cx + sw/2, h/2, 180, 360)
    c.restoreState()

def draw_compass_rose(c, cx, cy, r):
    """Silk Road compass star."""
    c.saveState()
    c.translate(cx, cy)
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.8)
    # Outer circles via 4-arc approximation
    def draw_circle_path(rr):
        """Draw a circle using cubic bezier arcs on canvas path."""
        p = c.beginPath()
        # Approximate circle with 4 cubic beziers
        k = 0.5522847498 * rr
        p.moveTo(rr, 0)
        p.curveTo(rr, k, k, rr, 0, rr)
        p.curveTo(-k, rr, -rr, k, -rr, 0)
        p.curveTo(-rr, -k, -k, -rr, 0, -rr)
        p.curveTo(k, -rr, rr, -k, rr, 0)
        return p
    c.drawPath(draw_circle_path(r), stroke=1, fill=0)
    c.drawPath(draw_circle_path(r*0.85), stroke=1, fill=0)
    
    for i in range(8):
        a = i * math.pi / 4
        dx, dy = math.cos(a), math.sin(a)
        ox, oy = dx * r * 0.95, dy * r * 0.95
        ix, iy = dx * r * 0.3, dy * r * 0.3
        if i % 2 == 0:
            c.setFillColor(GOLD)
            c.setStrokeColor(GOLD)
            p = c.beginPath()
            p.moveTo(ix, iy)
            p.lineTo(ox - dx*r*0.1, oy - dy*r*0.1)
            p.lineTo(ox, oy)
            p.lineTo(ix*0.5+dx*r*0.15*0.707, iy*0.5+dy*r*0.15*0.707)
            p.lineTo(ix, iy)
            c.drawPath(p, fill=1, stroke=1)
        else:
            c.setStrokeColor(LIGHT_GOLD)
            c.setLineWidth(0.5)
            p = c.beginPath()
            p.moveTo(ix*0.8, iy*0.8)
            p.lineTo(ox-dx*r*0.08, oy-dy*r*0.08)
            p.lineTo(ox, oy)
            p.lineTo(ix*1.2+dx*r*0.15, iy*1.2+dy*r*0.15)
            p.lineTo(ix*0.8, iy*0.8)
            c.drawPath(p, fill=0, stroke=1)
    # Center dot
    c.setFillColor(GOLD)
    c.setStrokeColor(GOLD)
    c.drawPath(draw_circle_path(r*0.08), fill=1, stroke=1)
    c.restoreState()

def draw_camel(c, x, y, s=1.0):
    """Simple camel silhouette using paths."""
    c.saveState()
    c.translate(x, y)
    c.scale(s, s)
    c.setFillColor(DARK_GOLD)
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.4)
    # Body + legs as a filled shape
    p = c.beginPath()
    p.moveTo(-30, 0)
    p.curveTo(-30, 16, -22, 22, -10, 24)
    p.curveTo(-5, 25, 5, 25, 12, 26)
    p.curveTo(18, 28, 24, 24, 28, 22)
    p.curveTo(32, 16, 32, 0, 32, 0)
    # Front legs (right)
    p.lineTo(30, -30)
    p.lineTo(27, -30)
    p.lineTo(27, -2)
    p.lineTo(10, -2)
    p.lineTo(10, -30)
    p.lineTo(7, -30)
    p.lineTo(7, -2)
    p.lineTo(0, 0)
    # Neck & head
    p.curveTo(-15, 10, -18, 18, -24, 22)
    p.curveTo(-28, 24, -34, 22, -36, 20)
    p.curveTo(-38, 17, -39, 14, -37, 12)
    p.curveTo(-35, 10, -28, 10, -26, 12)
    p.curveTo(-24, 10, -20, 2, -18, 0)
    p.lineTo(-30, 0)
    c.drawPath(p, fill=1, stroke=1)
    # Tail
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.8)
    p2 = c.beginPath()
    p2.moveTo(32, 18)
    p2.curveTo(38, 22, 40, 16, 36, 14)
    c.drawPath(p2, fill=0, stroke=1)
    c.restoreState()

def draw_dunes(c, w, h):
    """Desert dune silhouettes at bottom."""
    y0 = 22*mm
    c.setFillColor(DARK_GOLD)
    c.setStrokeColor(HexColor('#A0844A'))
    c.setLineWidth(0.3)
    p = c.beginPath()
    p.moveTo(0, y0)
    for i in range(41):
        x = i * w / 40
        wave = math.sin(i * 0.3) * 8*mm + math.sin(i * 0.7) * 4*mm
        p.lineTo(x, y0 + wave)
    p.lineTo(w, y0)
    p.lineTo(w, 0)
    p.lineTo(0, 0)
    p.lineTo(0, y0)
    c.drawPath(p, fill=1, stroke=1)

def draw_caravan(c, w, h):
    """Caravan of camels."""
    draw_camel(c, w*0.25, 40*mm, 0.8)
    draw_camel(c, w*0.48, 37*mm, 0.65)
    draw_camel(c, w*0.68, 35*mm, 0.5)

# ═══════════════════════════════════════════════════════
#  Page Decorators
# ═══════════════════════════════════════════════════════

def draw_page_border(c, w, h):
    """Decorative border for all content pages."""
    m = 14*mm
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.7)
    c.rect(m, m, w-2*m, h-2*m)
    c.setStrokeColor(LIGHT_GOLD)
    c.setLineWidth(0.3)
    c.rect(m+3*mm, m+3*mm, w-2*m-6*mm, h-2*m-6*mm)
    # Corners
    for cx, cy in [(m,m),(w-m,m),(m,h-m),(w-m,h-m)]:
        draw_corner_geom(c, cx, cy, 8*mm)

def draw_top_band(c, w):
    y = A4[1] - 11*mm
    m = 15*mm
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.6)
    c.line(m, y, w-m, y)
    draw_arabesque_band(c, m+5*mm, y+2*mm, w-2*m-10*mm)

def draw_bottom_band(c, w):
    y = 12*mm
    m = 15*mm
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.6)
    c.line(m, y, w-m, y)

def draw_section_hr(c, w, y):
    """Decorative horizontal rule with diamond center."""
    m = 15*mm
    lw = 50*mm
    cx = w / 2
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.6)
    c.line(cx - lw, y, cx - 10*mm, y)
    c.line(cx + 10*mm, y, cx + lw, y)
    c.setFillColor(GOLD)
    sz = 1.5*mm
    p = c.beginPath()
    p.moveTo(cx, y+sz)
    p.lineTo(cx+sz, y)
    p.lineTo(cx, y-sz)
    p.lineTo(cx-sz, y)
    p.lineTo(cx, y+sz)
    c.drawPath(p, fill=1, stroke=0)

# ═══════════════════════════════════════════════════════
#  Page Template Callbacks
# ═══════════════════════════════════════════════════════

def cover_page(c, doc):
    w, h = A4
    c.saveState()
    # Full deep blue background
    c.setFillColor(DEEP_BLUE)
    c.rect(0, 0, w, h, fill=1, stroke=0)
    # Border frames
    m = 12*mm
    c.setStrokeColor(GOLD)
    c.setLineWidth(1.5)
    c.rect(m, m, w-2*m, h-2*m)
    c.setStrokeColor(LIGHT_GOLD)
    c.setLineWidth(0.5)
    c.rect(m+4*mm, m+4*mm, w-2*m-8*mm, h-2*m-8*mm)
    # Corner geometrics
    for cx, cy in [(m,m),(w-m,m),(m,h-m),(w-m,h-m)]:
        draw_corner_geom(c, cx, cy, 14*mm)
    # Top decorative band
    draw_arabesque_band(c, m+10*mm, h-26*mm, w-2*(m+10*mm))
    # Compass rose
    draw_compass_rose(c, w/2, h-45*mm, 10*mm)
    # Title
    c.setFillColor(GOLD)
    c.setFont('msyh', 18)
    c.drawCentredString(w/2, h-72*mm, '新丝路跨境“实战”出海')
    c.setFont('msyh', 14)
    c.drawCentredString(w/2, h-78*mm, "第二套商业模式方案书")
    # Divider
    c.setStrokeColor(GOLD)
    c.setLineWidth(0.6)
    c.line(w/2-45*mm, h-83*mm, w/2+45*mm, h-83*mm)
    # Slogan
    c.setFont('msyh', 11)
    c.setFillColor(LIGHT_GOLD)
    c.drawCentredString(w/2, h-90*mm, "实战出海，交付为王")
    c.drawCentredString(w/2, h-96*mm, "——让每一次出海都有人陪跑到底")
    c.setFont('msyh', 10)
    c.setFillColor(HexColor('#D4C49A'))
    c.drawCentredString(w/2, h-104*mm, "以交付为基石的企业出海全链路服务平台")
    # Caravan & dunes
    draw_caravan(c, w, h)
    draw_dunes(c, w, h)
    # Bottom band
    draw_arabesque_band(c, m+10*mm, 26*mm, w-2*(m+10*mm))
    c.restoreState()

def content_page(c, doc):
    w, h = A4
    c.saveState()
    c.setFillColor(white)
    c.rect(0, 0, w, h, fill=1, stroke=0)
    draw_page_border(c, w, h)
    draw_top_band(c, w)
    draw_bottom_band(c, w)
    # Page number
    c.setFont('msyh', 8)
    c.setFillColor(DARK_GOLD)
    c.drawCentredString(w/2, 9*mm, f"— {doc.page} —")
    # Small corner geom accents
    draw_corner_geom(c, 20*mm, h-16*mm, 5*mm)
    draw_corner_geom(c, w-20*mm, h-16*mm, 5*mm)
    c.restoreState()


# ═══════════════════════════════════════════════════════
#  Custom Section Divider Flowable
# ═══════════════════════════════════════════════════════

class SilkRoadHR(Flowable):
    """Decorative horizontal divider flowable."""
    def __init__(self, width_pct=60, color=GOLD, thickness=0.8, space_before=2, space_after=4):
        Flowable.__init__(self)
        self.width_pct = width_pct
        self.color = color
        self.thickness = thickness
        self.space_before = space_before * mm
        self.space_after = space_after * mm

    def wrap(self, availWidth, availHeight):
        self._draw_w = availWidth * self.width_pct / 100.0
        h = self.thickness + 4*mm + self.space_before + self.space_after
        return (self._draw_w, h)

    def draw(self):
        c = self.canv
        c.saveState()
        cx = self._draw_w / 2.0
        lw = self._draw_w / 2 - 8*mm
        y = self.space_before + 2*mm
        c.setStrokeColor(self.color)
        c.setLineWidth(self.thickness)
        c.line(cx - lw, y, cx - 8*mm, y)
        c.line(cx + 8*mm, y, cx + lw, y)
        # Diamond
        c.setFillColor(self.color)
        sz = 2*mm
        p = c.beginPath()
        p.moveTo(cx, y+sz)
        p.lineTo(cx+sz, y)
        p.lineTo(cx, y-sz)
        p.lineTo(cx-sz, y)
        p.lineTo(cx, y+sz)
        c.drawPath(p, fill=1, stroke=0)
        c.restoreState()


# ═══════════════════════════════════════════════════════
#  Parse DOCX
# ═══════════════════════════════════════════════════════

def parse_docx(path):
    doc = Document(path)
    paras = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        sn = p.style.name
        if 'Heading 1' in sn:
            paras.append(('h1', t))
        elif 'Heading 2' in sn:
            paras.append(('h2', t))
        elif 'Heading 3' in sn:
            paras.append(('h3', t))
        else:
            paras.append(('body', t))
    tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        tables.append(rows)
    return paras, tables


# ═══════════════════════════════════════════════════════
#  Build Story Flowables
# ═══════════════════════════════════════════════════════

def build_story(paras, tables):
    story = []

    s_h1 = ParagraphStyle('H1', fontName='msyh', fontSize=16, textColor=DEEP_BLUE,
                          spaceBefore=10*mm, spaceAfter=5*mm, leading=24, alignment=TA_LEFT)
    s_h2 = ParagraphStyle('H2', fontName='msyh', fontSize=12, textColor=DARK_GOLD,
                          spaceBefore=6*mm, spaceAfter=3*mm, leading=18, alignment=TA_LEFT)
    s_h3 = ParagraphStyle('H3', fontName='msyh', fontSize=10.5, textColor=BRICK_RED,
                          spaceBefore=4*mm, spaceAfter=2*mm, leading=16, alignment=TA_LEFT)
    s_body = ParagraphStyle('Body', fontName='msyh', fontSize=9, textColor=INK_BLACK,
                            spaceBefore=0.8*mm, spaceAfter=1*mm, leading=14,
                            alignment=TA_LEFT)
    s_bullet = ParagraphStyle('Bullet', fontName='msyh', fontSize=9, textColor=INK_BLACK,
                              spaceBefore=0.3*mm, spaceAfter=0.3*mm, leading=14,
                              leftIndent=6*mm)

    th_style = ParagraphStyle('TH', fontName='msyh', fontSize=8, textColor=white,
                              alignment=TA_CENTER, leading=12)
    td_style = ParagraphStyle('TD', fontName='msyh', fontSize=7.5, textColor=INK_BLACK,
                              alignment=TA_LEFT, leading=12)

    # Cover page
    story.append(NextPageTemplate('cover'))
    story.append(PageBreak())
    story.append(NextPageTemplate('content'))

    # Build content
    for ptype, text in paras:
        if ptype == 'h1':
            story.append(Spacer(1, 2*mm))
            story.append(Paragraph(text, s_h1))
            story.append(SilkRoadHR())
        elif ptype == 'h2':
            story.append(Spacer(1, 1*mm))
            story.append(Paragraph(text, s_h2))
        elif ptype == 'h3':
            story.append(Paragraph(text, s_h3))
        else:
            if text[0:1] in ('•', '-', '*') or (text[0].isdigit() and '. ' in text[:5]):
                story.append(Paragraph(text, s_bullet))
            else:
                story.append(Paragraph(text, s_body))

    # Tables section
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph("附表及财务数据", s_h1))
    story.append(SilkRoadHR())

    for rows in tables:
        if not rows:
            continue
        n = len(rows[0])
        aw = 160*mm
        cw = [aw / n] * n

        wrapped = []
        for ri, row in enumerate(rows):
            wr = []
            for ci, cell in enumerate(row):
                cell = cell.strip() or " "
                wr.append(Paragraph(cell, th_style if ri == 0 else td_style))
            wrapped.append(wr)

        t = Table(wrapped, colWidths=cw, repeatRows=1)
        bg_alt = [HexColor('#FFFAF0'), HexColor('#FFF5E0')]
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), DEEP_BLUE),
            ('TEXTCOLOR', (0,0), (-1,0), white),
            ('ALIGN', (0,0), (-1,0), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('FONTNAME', (0,0), (-1,-1), 'msyh'),
            ('FONTSIZE', (0,0), (-1,-1), 7.5),
            ('GRID', (0,0), (-1,-1), 0.3, GOLD),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 3),
            ('RIGHTPADDING', (0,0), (-1,-1), 3),
        ]))
        # Alternating row colors for non-header rows
        for ri in range(1, len(wrapped)):
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,ri), (-1,ri), bg_alt[ri % 2])
            ]))

        story.append(Spacer(1, 4*mm))
        story.append(t)
        story.append(Spacer(1, 3*mm))

    return story


# ═══════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════

def generate(docx_path, pdf_path):
    print(f"Reading: {docx_path}")
    paras, tables = parse_docx(docx_path)
    print(f"  Paragraphs: {len(paras)}, Tables: {len(tables)}")

    print("Registering fonts...")
    register_fonts()

    print("Building document...")
    doc = BaseDocTemplate(pdf_path, pagesize=A4,
                          leftMargin=22*mm, rightMargin=22*mm,
                          topMargin=20*mm, bottomMargin=18*mm)

    fc = Frame(22*mm, 18*mm, A4[0]-44*mm, A4[1]-36*mm, id='cframe')

    doc.addPageTemplates([
        PageTemplate(id='cover', frames=Frame(0, 0, A4[0], A4[1], id='cover_frame'),
                     onPage=cover_page, pagesize=A4),
        PageTemplate(id='content', frames=fc,
                     onPage=content_page, pagesize=A4),
    ])

    story = build_story(paras, tables)
    print("Generating PDF...")
    doc.build(story)

    if os.path.exists(pdf_path):
        size = os.path.getsize(pdf_path) / 1024
        print(f"SUCCESS! PDF: {pdf_path} ({size:.1f} KB)")
    else:
        print("ERROR: File not created")
        return False
    return True

if __name__ == '__main__':
    src = r'C:\Users\ASDCF\Desktop\新丝路跨境_实战出海_第二套商业模式方案书.docx'
    dst = r'C:\Users\ASDCF\Desktop\新丝路跨境_实战出海_第二套商业模式方案书.pdf'
    generate(src, dst)
