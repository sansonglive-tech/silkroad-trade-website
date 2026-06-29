#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
新丝路·企业主会 商业计划书 Word文档生成
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os
from datetime import datetime

def create_xinsilu_document():
    """创建新丝路商业计划书"""
    doc = Document()
    
    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # ========== 标题样式 ==========
    # 主标题样式
    title_style = doc.styles['Title']
    title_style.font.name = '微软雅黑'
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    # 标题1样式
    heading1 = doc.styles['Heading 1']
    heading1.font.name = '微软雅黑'
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    
    # 标题2样式
    heading2 = doc.styles['Heading 2']
    heading2.font.name = '微软雅黑'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)
    
    # 标题3样式
    heading3 = doc.styles['Heading 3']
    heading3.font.name = '微软雅黑'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0x0d, 0x6e, 0xad)
    
    # 正文样式
    normal = doc.styles['Normal']
    normal.font.name = '微软雅黑'
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)
    
    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('新丝路·企业主会')
    run.font.name = '微软雅黑'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x5f, 0x9a)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('基于IP存量的全链路出海服务平台')
    run.font.name = '微软雅黑'
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    
    doc.add_paragraph()
    
    # 版本信息
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('商业计划书')
    run.font.name = '微软雅黑'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 标签
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run('99元引流课锚点版 · 完整方案')
    run.font.name = '微软雅黑'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
    run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 日期
    date_str = datetime.now().strftime('%Y年%m月')
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(date_str)
    run.font.name = '微软雅黑'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_page_break()
    
    # ========== 一、执行摘要 ==========
    doc.add_heading('一、执行摘要', 1)
    
    exec_summary = """新丝路·企业主会是一个以创始人20年外贸实战经验 + 几十万粉丝IP存量为核心资产，借鉴"山海图"深度本地化服务模式，为国内中小制造企业提供"认知升级→技能培训→线下深交→一站式出海服务"的全链路解决方案。"""
    doc.add_paragraph(exec_summary)
    
    doc.add_heading('核心商业模式', 2)
    business_model = """以99元引流课为信任钩子，通过线上内容矩阵（短视频/直播/白皮书）和私域分层运营，将粉丝精准转化为付费用户；再通过进阶课程、付费社群、线下会销（大课/私董会/峰会/海外考察）实现深度转化；最终以海外公司注册、财税合规、知识产权、供应链金融等平台服务完成长期价值绑定，形成"前端让利获客、后端服务盈利"的可持续生态。"""
    doc.add_paragraph(business_model)
    
    doc.add_heading('差异化壁垒', 2)
    barriers = [
        'IP信任前置（几十万粉丝已建立认知）',
        '20年跨国外贸实战经验（内容不可复制）',
        '山海图式本地化服务网络（轻资产合作+海外资源）',
        '99元引流课→高客单价服务的平滑阶梯'
    ]
    for item in barriers:
        p = doc.add_paragraph(f'• {item}')
    
    doc.add_paragraph()
    
    # 目标
    target = doc.add_paragraph()
    run = target.add_run('3年成熟期营收目标：')
    run.font.bold = True
    run = target.add_run('5,000万+，净利率20-25%')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    
    doc.add_page_break()
    
    # ========== 二、市场机遇与差异化定位 ==========
    doc.add_heading('二、市场机遇与差异化定位', 1)
    
    doc.add_heading('2.1 行业背景', 2)
    industry_bg = [
        '国内制造业内卷加剧，出海成为企业破局核心路径',
        '跨境电商与外贸服务市场规模持续增长，2026年外贸培训类目预计达620亿元',
        '国家政策支持（《对外贸易法》修订、跨境电商出口海外仓"离境即退税"等）',
        '知识付费进入3.0时代，AI驱动内容生产，用户更看重"结果导向"的实战交付'
    ]
    for item in industry_bg:
        doc.add_paragraph(f'• {item}')
    
    doc.add_heading('2.2 "新丝路"品牌内涵', 2)
    
    brand_quote = doc.add_paragraph()
    run = brand_quote.add_run('新丝路——传承千年丝绸之路的开拓精神，在新时代为中国制造搭建通往全球的桥梁。')
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1a, 0x5f, 0x9a)
    
    brand_meaning = [
        '新：新理念、新方法、AI新工具、全球新市场',
        '丝路：联通、合作、信任、悠远的历史底蕴',
        '品牌口号：新丝路——让中国制造走得更远'
    ]
    for item in brand_meaning:
        doc.add_paragraph(f'• {item}')
    
    doc.add_heading('2.3 客户定位', 2)
    customer = [
        '客户定位：年营收500万-2亿元的国内制造企业主/工厂负责人',
        '服务定位：比传统咨询更懂工厂，比普通服务平台更系统化',
        '价值主张：用20年实战经验 + 全球本地化网络，帮助企业降低出海试错成本，缩短盈利周期'
    ]
    for item in customer:
        doc.add_paragraph(f'• {item}')
    
    doc.add_page_break()
    
    # ========== 三、商业模式总览 ==========
    doc.add_heading('三、商业模式总览', 1)
    
    intro = doc.add_paragraph('以IP资产为核心，通过四层漏斗实现用户全生命周期价值最大化：')
    
    layers = [
        ('第一层：99元引流课', '短视频/直播/白皮书（免费）→ 日均触达10万+；99元系列课；私域分层运营'),
        ('第二层：进阶课程', '进阶线上课（499-1,999元）；"新丝路研习社"年度会员（1,980-9,800元）'),
        ('第三层：线下会销', '出海大课城市巡回（3,800-5,800元/2天）；新丝路私董会闭门（12,800-29,800元/期）；新丝路年度出海峰会；海外考察团（38,000-88,000元/人）'),
        ('第四层：出海平台服务', '海外公司注册（5,000-15,000元/单）；年度财税/合规托管（6,000-30,000元/年）；知识产权代理（3,000-20,000元/单）；供应链金融/物流佣金（交易额2-5%）')
    ]
    
    for i, (title, content) in enumerate(layers, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{title}：')
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        p.add_run(content)
    
    doc.add_paragraph()
    
    profit_summary = doc.add_paragraph()
    run = profit_summary.add_run('盈利闭环：')
    run.font.bold = True
    run = profit_summary.add_run('99元让利获客 → 后端服务赚取长期价值（LTV 2,400元+）')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)
    
    doc.add_page_break()
    
    # ========== 四、第一层：99元引流课 ==========
    doc.add_heading('四、第一层：99元引流课为核心的线上获客体系', 1)
    
    doc.add_heading('4.1 产品设计', 2)
    
    courses = [
        ('《外贸AI获客实战课》', 'AI提示词生成开发信、客户画像、邮件自动化', '5节×15分钟', '99元', 'AI提示词库+SOP模板'),
        ('《LinkedIn外贸7天训练营》', '账号搭建、内容策略、InMail技巧、人脉扩展', '7天打卡', '99元', '话术模板+每日任务'),
        ('《外贸开发信从0到1》', '标题、正文、跟进策略、A/B测试', '4节×20分钟', '99元', '20个高回复率模板')
    ]
    
    for name, content, duration, price, value in courses:
        p = doc.add_paragraph()
        run = p.add_run(f'• {name}')
        run.font.bold = True
        p.add_run(f' | 内容：{content}')
        p.add_run(f' | 课时：{duration}')
        p.add_run(f' | 定价：{price}')
        p.add_run(f' | 附加：{value}')
    
    doc.add_heading('4.2 设计原则', 2)
    principles = [
        '超值感：交付价值感远高于99元',
        '低门槛：每节课15-20分钟，碎片化学习',
        '即时反馈：第1节课即给可操作技巧',
        '埋钩子：课程中预告进阶内容，引导转化'
    ]
    for item in principles:
        doc.add_paragraph(f'• {item}')
    
    doc.add_heading('4.3 私域分层', 2)
    segments = [
        ('企业主群', '年营收500万+，决策者', '宏观趋势、市场分析、风险合规', '私董会、平台服务'),
        ('外贸从业者群', '打工/创业初期，执行层', '战术技巧、工具使用、案例拆解', '进阶课、会员'),
        ('小白群', '新手/学生，兴趣导向', '基础概念、职业发展', '99元课、入门课')
    ]
    
    for seg, features, content_strat, goal in segments:
        p = doc.add_paragraph()
        run = p.add_run(f'• {seg}：')
        run.font.bold = True
        p.add_run(f'{features} → {content_strat} → 转化目标：{goal}')
    
    doc.add_page_break()
    
    # ========== 五、第二层：进阶课程与付费社群 ==========
    doc.add_heading('五、第二层：进阶课程与付费社群', 1)
    
    doc.add_heading('5.1 进阶线上课', 2)
    advanced_courses = [
        ('《独立站+谷歌SEO全链路》', '建站、关键词、外链、转化率优化', '999元'),
        ('《跨境电商合规与海外公司架构》', '美国/欧洲/东南亚公司选择、VAT、商标', '1,499元'),
        ('《20年外贸老兵教你拿下大客户》', '客户开发、谈判、关系维护、案例库', '1,999元')
    ]
    
    for name, content, price in advanced_courses:
        p = doc.add_paragraph()
        run = p.add_run(f'• {name}')
        run.font.bold = True
        p.add_run(f'：{content} | 定价：{price}')
    
    doc.add_heading('5.2 付费社群', 2)
    communities = [
        ('新丝路研习社', '每月闭门分享、资源对接、课程8折、线下优先', '1,980元/年', '500人'),
        ('新丝路私董会（审核制）', '全年4次线下私享会、1v1顾问、海外考察优先、创始人群', '9,800元/年', '100人')
    ]
    
    for name, rights, price, target in communities:
        p = doc.add_paragraph()
        run = p.add_run(f'• {name}')
        run.font.bold = True
        p.add_run(f' | 权益：{rights}')
        p.add_run(f' | 年费：{price}')
        p.add_run(f' | 目标：{target}')
    
    doc.add_page_break()
    
    # ========== 六、第三层：线下会销矩阵 ==========
    doc.add_heading('六、第三层：线下会销矩阵', 1)
    
    offline_products = [
        ('出海大课', '2天集中授课+案例拆解+晚宴', '100-300人', '3,800-5,800元', '每月1-2城', '批量转化进阶用户'),
        ('新丝路私董会', '闭门研讨+1v1诊断+圆桌', '10-20人', '12,800-29,800元/期', '每季度1期', '高净值客户深度绑定'),
        ('新丝路年度出海峰会', '主题演讲+圆桌+资源展+颁奖', '500-1,000人', '980-2,980元/人', '每年1次', '行业影响力+品牌IP'),
        ('海外考察团', '7-12天，参访+政府对接+市场调研', '15-30人', '38,000-88,000元/人', '每半年1次', '实地落地+高客单价')
    ]
    
    for name, form, scale, price, freq, goal in offline_products:
        p = doc.add_paragraph()
        run = p.add_run(f'• {name}')
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        p.add_run(f'\n  形式：{form} | 规模：{scale} | 定价：{price} | 频率：{freq} | 目标：{goal}')
        doc.add_paragraph()
    
    doc.add_heading('城市巡回计划', 2)
    cities = '第一年覆盖城市：深圳、广州、宁波、苏州、青岛、成都、重庆、武汉（8城）'
    doc.add_paragraph(cities)
    
    partners = '合作方：本地跨境电商协会、产业带商会、外贸综试区'
    doc.add_paragraph(partners)
    
    doc.add_page_break()
    
    # ========== 七、第四层：出海一体化服务平台 ==========
    doc.add_heading('七、第四层：出海一体化服务平台', 1)
    
    services = [
        ('海外公司注册', '美国LLC、英国、新加坡、香港公司', '5,000-15,000元/单 + 年审', '自营+本地律所分成'),
        ('财税合规', 'VAT注册申报、年度审计、税务筹划', '6,000-30,000元/年', '合作会计所'),
        ('知识产权', '美国/欧盟商标、专利申请', '3,000-20,000元/单', '合作知识产权代理'),
        ('银行开户', '远程开立海外公司账户', '2,000-5,000元/单', '合作金融机构'),
        ('产品准入', 'CE/FCC/ROHS等认证咨询', '按项目报价', '合作实验室'),
        ('供应链金融', '出口信用保险、订单融资', '佣金2-5%', '合作银行/保理'),
        ('国际物流', '海运/空运/海外仓对接', '佣金3-8%', '合作物流平台')
    ]
    
    for name, content, price, partner in services:
        p = doc.add_paragraph()
        run = p.add_run(f'• {name}')
        run.font.bold = True
        p.add_run(f'：{content} | 收费：{price} | 合作：{partner}')
    
    doc.add_page_break()
    
    # ========== 八、盈利测算 ==========
    doc.add_heading('八、盈利测算', 1)
    
    doc.add_heading('8.1 三年收入预测', 2)
    
    revenue_data = [
        ('L1引流课', '99元', '99万', '198万', '297万'),
        ('L2进阶课', '899元', '45万', '135万', '270万'),
        ('L3付费社群', '4,800元', '48万', '144万', '384万'),
        ('L4线下会销', '15,000元', '150万', '450万', '900万'),
        ('L5平台服务', '8,000元/企', '80万', '400万', '1,200万'),
        ('L6 B端合作', '5万/单', '50万', '150万', '300万'),
        ('L7生态衍生', '-', '50万', '300万', '2,000万+')
    ]
    
    # 表头
    header = doc.add_paragraph()
    run = header.add_run('盈利层'.ljust(15))
    run.font.bold = True
    header.add_run('客单价'.ljust(12))
    header.add_run('第1年'.ljust(12))
    header.add_run('第2年'.ljust(12))
    header.add_run('第3年'.ljust(12))
    
    for layer, price, y1, y2, y3 in revenue_data:
        p = doc.add_paragraph()
        run = p.add_run(layer.ljust(15))
        run.font.bold = True
        p.add_run(price.ljust(12))
        p.add_run(y1.ljust(12))
        p.add_run(y2.ljust(12))
        p.add_run(y3.ljust(12))
    
    doc.add_paragraph()
    total = doc.add_paragraph()
    run = total.add_run('合计：')
    run.font.bold = True
    run.font.size = Pt(12)
    run = total.add_run('第1年 522万 | 第2年 1,777万 | 第3年 5,351万+')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)
    
    doc.add_heading('8.2 核心财务指标', 2)
    finance_metrics = [
        '单个99元用户LTV：约2,400元',
        '获客成本（CAC）：可控在300元以内',
        'LTV/CAC：>8（健康）',
        '整体毛利率：30-35%',
        '净利率（成熟期）：20-25%',
        '投资回收期：12-18个月'
    ]
    for item in finance_metrics:
        doc.add_paragraph(f'• {item}')
    
    doc.add_page_break()
    
    # ========== 九、实施路线图 ==========
    doc.add_heading('九、实施路线图', 1)
    
    phases = [
        ('阶段一：品牌奠基期（0-6个月）', [
            '粉丝分层梳理、99元课大纲设计、AI内容工具搭建',
            '录制99元引流课、上线小鹅通/私域、发布首期白皮书',
            '99元课正式推广，跑通"购买-入群-训练营-转化"全流程',
            '举办首场线下大课（100人以下），验证会销模型；启动海外合作渠道'
        ]),
        ('阶段二：规模化增长期（6-18个月）', [
            '进阶课程上线、新丝路研习社会员启动',
            '私董会首期举办、年度峰会筹备、平台服务V1.0上线',
            '城市巡回扩至8城、海外考察团首发',
            '平台服务客户累计200+；私域企业主粉丝达5万+'
        ]),
        ('阶段三：生态成熟期（18-36个月）', [
            '数字化服务平台上线、开放第三方服务商入驻',
            '生态衍生收入规模化（佣金/金融/撮合）',
            '打造出海生态"智库"，覆盖30+产业带城市'
        ])
    ]
    
    for phase_title, tasks in phases:
        doc.add_heading(phase_title, 2)
        for task in tasks:
            doc.add_paragraph(f'• {task}')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 十、风险与防范 ==========
    doc.add_heading('十、风险与防范', 1)
    
    risks = [
        ('IP信任透支', '交付不及预期，粉丝流失', '坚持实战案例+真实数据，承诺可控，超值交付'),
        ('线下活动冷场', '报名人数不足或体验差', '先用99元课用户邀约，保证基础人数；打磨标准化SOP'),
        ('服务交付瓶颈', '后端服务能力跟不上销售', '前期采用合作分成模式，自营核心服务，控制接单节奏'),
        ('竞争对手模仿', '其他IP复制类似模式', '壁垒在于20年经验+海外资源，持续强化内容深度'),
        ('跨境合规风险', '涉及海外法律税务', '只与持牌机构合作，购买合规保险，客户告知书明确责任边界')
    ]
    
    for risk, desc, solution in risks:
        p = doc.add_paragraph()
        run = p.add_run(f'• {risk}')
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        p.add_run(f'\n  风险：{desc}\n  防范：{solution}')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 十一、核心竞争力总结 ==========
    doc.add_heading('十一、核心竞争力总结', 1)
    
    core_advantages = [
        ('IP认知壁垒', '几十万粉丝的信任 + 20年外贸实战 = 出海知识领域的第一KOL心智'),
        ('私域企业主社群', '粉丝中沉淀了大量国内制造企业决策者，形成精准获客池'),
        ('海外本地化网络', '20年积累的律师、会计师、物流、政府资源，快速搭建服务生态'),
        ('99元→高客单价平滑阶梯', '产品线设计科学，用户自然向上流动，拒绝"割韭菜"'),
        ('结果导向的闭环', '不只是卖课，而是真正帮助企业注册公司、合规运营、拿到订单'),
        ('丝路文化底蕴+政策高度', '品牌名自带"一带一路"红利，更容易获得政府与商会资源支持')
    ]
    
    for i, (title, desc) in enumerate(core_advantages, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title}')
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        p.add_run(f'\n   {desc}')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 十二、启动行动清单 ==========
    doc.add_heading('十二、启动行动清单（第一个月）', 1)
    
    action_items = [
        ('第1周', '梳理粉丝画像，统计企业主占比、行业分布、地域分布；整理20年外贸案例库'),
        ('第2周', '用AI辅助产出第一版99元课大纲和3节样片；注册"新丝路"品牌商标和公众号/视频号'),
        ('第3周', '发布首条品牌宣发视频，私域预热99元课；联系3个地方商会洽谈合作'),
        ('第4周', '99元课正式上架小鹅通，启动第一轮推广（粉丝群+朋友圈+直播）；确定首场线下大城市（深圳）')
    ]
    
    for week, action in action_items:
        p = doc.add_paragraph()
        run = p.add_run(f'• {week}：')
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
        p.add_run(action)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========== 结束语 ==========
    ending = doc.add_paragraph()
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ending.add_run('新丝路——让中国制造走得更远。🌏')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x5f, 0x9a)
    
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('本商业模型为创始人20年经验与IP势能的系统化落地，欢迎进一步探讨细节模块的深化设计。')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # ========== 保存文档 ==========
    output_path = r'E:\郑州录音\新丝路·企业主会_商业计划书.docx'
    doc.save(output_path)
    
    size_kb = os.path.getsize(output_path) / 1024
    
    print(f"[SUCCESS] 文档已保存: {output_path}")
    print(f"[INFO] 文件大小: {size_kb:.1f} KB")
    print(f"[INFO] 包含12个章节，完整商业计划书结构")
    
    return output_path

if __name__ == '__main__':
    print("="*60)
    print("正在生成新丝路·企业主会商业计划书 Word文档...")
    print("="*60)
    
    output = create_xinsilu_document()
    
    if output:
        print("\n[SUCCESS] Word文档生成完成！")
        print(f"[INFO] 文件位置: {output}")
