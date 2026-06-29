# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = docx.Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)

def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_colored_run(para, text, size=11, bold=False, color=None, italic=False):
    run = para.add_run(text)
    set_font(run, size, bold, color, italic)
    return run

# ── 封面 ──
def make_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    add_colored_run(p, '外贸全链条培训记录', size=36, bold=True, color=(0x06, 0x5A, 0x82))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_colored_run(p2, '身份选择 · 外汇管理 · 海关报关 · 税务合规 · 品牌运营',
                    size=16, color=(0x1C, 0x72, 0x93))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    add_colored_run(p3, '内部培训资料', size=14, color=(0x64, 0x74, 0x8B))

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(6)
    add_colored_run(p4, '培训部  |  2025年版', size=11, color=(0x64, 0x74, 0x8B))

make_cover(doc)
doc.add_page_break()

# ── 目录 ──
def make_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_colored_run(p, '目  录', size=24, bold=True, color=(0x06, 0x5A, 0x82))

    toc_items = [
        ('01', '外贸人才库与三大关键要素'),
        ('02', '做外贸的三种身份（上）：个人身份'),
        ('03', '做外贸的三种身份（下）：1039模式与公司'),
        ('04', '三部门监管框架'),
        ('05', '外汇管理与收结汇实务'),
        ('06', '海关报关全流程'),
        ('07', '税务合规与退税实务'),
        ('08', '品牌运营官：客户开发方法论'),
        ('09', '人财物关系与思维转型'),
        ('10', '总结与30天行动清单'),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        add_colored_run(p, num + '  ', size=14, bold=True, color=(0x1C, 0x72, 0x93))
        add_colored_run(p, title, size=14, color=(0x1E, 0x29, 0x3B))

make_toc(doc)
doc.add_page_break()

# ── 章节标题 ──
def add_chapter(doc, num, title, subtitle=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    add_colored_run(p, f'第{num}章  {title}', size=20, bold=True, color=(0x06, 0x5A, 0x82))
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        add_colored_run(p2, subtitle, size=13, color=(0x64, 0x74, 0x8B), italic=True)
    # 分隔线（用下划线段落模拟）
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(8)
    run = p3.add_run('―' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1C, 0x72, 0x93)

# ── 小节标题 + 内容 ──
def add_section(doc, title, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_colored_run(p, title, size=14, bold=True, color=(0x06, 0x5A, 0x82))

    for item in items:
        p2 = doc.add_paragraph(style='List Bullet')
        add_colored_run(p2, item, size=11, color=(0x1E, 0x29, 0x3B))
        p2.paragraph_format.space_after = Pt(4)

# ── 提示框（行动建议）──
def add_tip_box(doc, icon_text, title, items, bg_hex='E8F4F8', border_hex='1C7293'):
    # 用表格模拟提示框
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.width = Cm(16)

    # 标题行
    tp = cell.add_paragraph()
    add_colored_run(tp, icon_text + '  ' + title, size=13, bold=True, color=(0x06, 0x5A, 0x82))

    # 内容
    for item in items:
        ip = cell.add_paragraph(style='List Bullet')
        add_colored_run(ip, item, size=11, color=(0x1E, 0x29, 0x3B))
        ip.paragraph_format.space_after = Pt(3)

    cell.paragraphs[0].paragraph_format.space_after = Pt(6)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════
# 第1章
# ══════════════════════════════════════
add_chapter(doc, '1', '外贸人才库与三大关键要素', '人才是外贸业务的基础设施')

add_section(doc, '外贸人才库的战略地位', [
    '人才库是外贸业务的基础设施，不是可选项',
    '忽视人才库的三大风险：财务流失、客户丢失、产品断供',
    '建立系统化人才筛选和培养机制',
])
add_section(doc, '培训目标', [
    '从「懂一点」到「全链条可控」',
    '对照自身业务，识别薄弱环节并制定改进计划',
    '掌握身份选择、外汇管理、海关报关、税务合规四大核心能力',
])
add_tip_box(doc, '💡', '行动建议', [
    '对照自身业务，识别薄弱环节并制定改进计划',
    '每季度评估一次人才库，确保关键岗位有备份人选',
])

doc.add_page_break()

# ══════════════════════════════════════
# 第2章
# ══════════════════════════════════════
add_chapter(doc, '2', '做外贸的三种身份（上）', '个人身份：利弊与定位')

add_section(doc, '个人身份特征', [
    '不能直接收外汇，无独立采购主体',
    '无退税资质，收入依赖提成',
    '本质角色：业务中间人',
])
add_section(doc, '核心风险', [
    '客户关系不受控 → 客户流失',
    '货源不受控 → 供应链断裂',
    '资金不受控 → 提成被克扣',
])
add_tip_box(doc, '⚠', '适用场景与建议', [
    '适用：刚入行试水、学习行业经验阶段',
    '建议：尽快升级到个体户或公司身份，掌握主动权',
    '警惕：「黑心老板」克扣提成，客户被老板直接接管',
], bg_hex='FFF4E6', border_hex='F59E0B')

doc.add_page_break()

# ══════════════════════════════════════
# 第3章
# ══════════════════════════════════════
add_chapter(doc, '3', '做外贸的三种身份（下）', '1039模式与公司身份')

add_section(doc, '1039市场采购贸易模式（义乌模式）', [
    '限定主体、限定区域、限定收款银行',
    '单票不超过15万美金，无需退税，即买即走',
    '不承担报关义务，多种灵活收款方式',
    '适用区域：义乌、白沟箱包、临沂等产业带',
    '适合：超能卷的产业带、农村作坊式生产',
])
add_section(doc, '公司身份：正规军路线', [
    '必须具备外贸资质并完成海关/外汇/税务备案',
    '通过SWIFT系统收汇，拥有对公账户',
    '可享受出口退税（最高13%）——这是纯利润',
    '人财物全可控，客户关系稳固，可规模化发展',
])
add_tip_box(doc, '💰', '退税价值计算', [
    '年出口额1000万 → 退税可达130万',
    '这是正规外贸企业的重要利润来源，务必重视',
], bg_hex='ECFDF5', border_hex='065F46')

doc.add_page_break()

# ══════════════════════════════════════
# 第4章
# ══════════════════════════════════════
add_chapter(doc, '4', '三部门监管框架', '外贸监管体系全景')

add_section(doc, '三大监管部门', [
    '外汇管理局（管钱）：SWIFT收汇审核、结汇合规管理、冻卡风险控制',
    '海关（管货）：报关单审核、查验（抽检/必检）、放行与边检',
    '税务/公安（管账与合规）：退税资质审核、账务合规检查、非法经营打击',
])
add_section(doc, '监管红线（严禁触碰）', [
    '公安介入场景：洗钱、非法经营、冻卡',
    '冻卡高发原因：地下钱庄、买单出口——碰都不要碰',
    '合规是底线，不是选择',
])
add_tip_box(doc, '⚠', '风控建议', [
    '定期自查三个维度（外汇/海关/税务）的合规状态',
    '建立风控清单，确保所有交易留痕可查',
], bg_hex='FEE2E2', border_hex='F96167')

doc.add_page_break()

# ══════════════════════════════════════
# 第5章
# ══════════════════════════════════════
add_chapter(doc, '5', '外汇管理与收结汇实务', 'SWIFT国际结算系统')

add_section(doc, 'SWIFT系统流程', [
    '客户TT电汇 → 水单/报文 → 外汇管理局审核 → 结汇至国内账户',
    '以美元为锚定货币，全球银行间标准化通信网络',
    '每一笔外汇流入都有迹可循，合规记录至关重要',
])
add_section(doc, '银行选择建议（服务评级）', [
    '招商银行：⭐⭐⭐⭐⭐ 最佳 —— 服务态度最好，效率最高，强烈推荐',
    '中信银行：⭐⭐⭐⭐ 优秀 —— 服务良好，效率较高',
    '工商银行：⭐⭐⭐ 良好 —— 服务不错，效率中等',
    '中国银行：⭐⭐⭐ 良好 —— 传统外贸银行，经验丰富',
    '农业银行：⭐ 较差 —— 效率偏低，流程慢，需预留更多时间',
])
add_tip_box(doc, '💡', '开户建议', [
    '优先选择招商银行开立外汇账户',
    '熟悉水单核验流程，确保收款凭证真实有效',
])

doc.add_page_break()

# ══════════════════════════════════════
# 第6章
# ══════════════════════════════════════
add_chapter(doc, '6', '海关报关全流程', '从申报到放行')

add_section(doc, '报关五大步骤', [
    '01 如实申报：品名、规格、单价、总价、HS编码，报关单信息必须真实准确',
    '02 海关查验：随机抽查 + 必检品类（医药、食品等）',
    '03 三流合一核对：货物流、资金流、单证流三者必须一致',
    '04 放行 + 边检：放行单 + 边检查违禁品（毒品、走私物品）',
    '05 目的国报关：产品认证、原产地签证，符合目的国进口合规要求',
])
add_tip_box(doc, '💡', '实操建议', [
    '提前确认产品归类，准备好完整单证',
    '如实申报是底线，虚报瞒报面临行政处罚甚至刑事追责',
])

doc.add_page_break()

# ══════════════════════════════════════
# 第7章
# ══════════════════════════════════════
add_chapter(doc, '7', '税务合规与退税实务', '退税是重要利润来源')

add_section(doc, '出口退税机制', [
    '出口退税最高可达13%，是正规外贸企业的重要利润来源',
    '前提：账务合规，发票、报关单、收汇凭证三者匹配',
    '「平汇、平账」原则：外汇金额与报关金额必须一致',
])
add_section(doc, '常见陷阱（严禁！）', [
    '买单出口：短期省钱，实为冻卡和处罚埋雷',
    '地下钱庄结汇：直接触发公安打击',
    '涉黑资金出海：风险极大，坚决不碰',
])
add_tip_box(doc, '✅', '行动建议', [
    '宁可合规成本高一些，也不要在税务和资金通道上走捷径',
    '定期自查三个维度（外汇/海关/税务）的合规状态',
    '建立风控清单，确保所有交易留痕可查',
], bg_hex='ECFDF5', border_hex='065F46')

doc.add_page_break()

# ══════════════════════════════════════
# 第8章
# ══════════════════════════════════════
add_chapter(doc, '8', '品牌运营官：客户开发方法论', '列名单 + 扔钩子')

add_section(doc, '列名单：锁定精准目标客户', [
    '核心动作：梳理并锁定20个目标客户',
    '客户画像标准：有需求、有预算、可触达、有决策权',
    '名单质量决定后续转化效率',
    '每周更新名单，持续做分层和优先级排序',
])
add_section(doc, '扔钩子：激发客户好奇心', [
    '钩子设计原则：制造信息差，引发客户主动询问',
    '不卖产品，卖「价值预期」和「合作想象」',
    '好钩子：案例故事、行业数据、政策红利',
    '为每个客户定制差异化沟通切入点',
])
add_tip_box(doc, '🎣', '钩子示例（中医出海场景）', [
    '「我们已经帮XX省的中医馆在东南亚落地了3家，单月营收破50万，您有兴趣了解具体模式吗？」',
])

doc.add_page_break()

# ══════════════════════════════════════
# 第9章
# ══════════════════════════════════════
add_chapter(doc, '9', '人财物关系与思维转型', '从内贸到外贸')

add_section(doc, '身份决定人财物格局', [
    '人：你是谁？代表谁？—— 决定了客户对你的信任基础',
    '财：钱走谁的账？—— 决定了利润空间和合规要求',
    '物：货从哪来？—— 决定了供应链控制力和议价能力',
])
add_section(doc, '内贸思维 vs 外贸思维', [
    '内贸：关系驱动，现货现款，灵活变通',
    '外贸：合规驱动，信用证/TT结算，流程标准化',
    '思维转变的关键：把「合规」视为竞争力，而非负担',
])
add_tip_box(doc, '💡', '转型建议', [
    '梳理现有业务流程，标注哪些需要按外贸标准改造',
    '把合规流程当作企业核心竞争力来建设',
])

doc.add_page_break()

# ══════════════════════════════════════
# 第10章
# ══════════════════════════════════════
add_chapter(doc, '10', '总结与30天行动清单', '从培训到执行')

add_section(doc, '培训核心要点回顾', [
    '身份选择决定业务天花板：个人 < 1039个体户 < 公司（正规军）',
    '三部门监管（外汇/海关/税务）是外贸运营的硬骨架，缺一不可',
    '合规是底线，不是选择——冻卡、处罚的风险远大于合规成本',
    '品牌运营官的核心：列名单 + 扔钩子，系统化开发客户',
    '人财物关系清晰，才能在外贸业务中掌握主动权',
])

doc.add_page_break()

# ── 30天行动清单 ──
add_chapter(doc, '附录', '30天行动清单', '可执行的具体事项')

checklist = [
    '确认当前身份是否匹配业务规模，制定升级计划',
    '开立合规外汇账户（推荐：招商银行）',
    '梳理报关单证模板，确保「三流合一」',
    '列出20个目标客户名单，设计差异化钩子',
    '完成税务自查，确认退税资质和流程通畅',
    '建立「人财物」关系图谱，明确权责边界',
]
for i, item in enumerate(checklist, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_colored_run(p, f'□  {i}. ', size=12, bold=True, color=(0x06, 0x5A, 0x82))
    add_colored_run(p, item, size=12, color=(0x1E, 0x29, 0x3B))

# ── 结束页 ──
doc.add_page_break()
ep = doc.add_paragraph()
ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
ep.paragraph_format.space_before = Pt(100)
add_colored_run(ep, '谢谢聆听', size=36, bold=True, color=(0x06, 0x5A, 0x82))

ep2 = doc.add_paragraph()
ep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_run(ep2, '合规为基  ·  客户为王  ·  执行致胜', size=14, color=(0x64, 0x74, 0x8B))

# 保存
out = r'E:\郑州录音\外贸全链条培训记录_文字版.docx'
doc.save(out)
print('Done: ' + out)
