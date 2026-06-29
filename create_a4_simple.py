#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from datetime import datetime

def clean_text(text):
    """清理语音识别噪音"""
    text = re.sub(r'$$\d+\.\d+s$$', '', text)
    text = re.sub(r'([。！？])\1+', r'\1', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def read_all_txt():
    """读取所有txt文件"""
    files = [
        r'E:\郑州录音\2.txt',
        r'E:\郑州录音\3.txt',
        r'E:\郑州录音\4.txt',
        r'E:\郑州录音\5.txt',
        r'E:\郑州录音\6.txt',
    ]
    
    all_text = []
    for fpath in files:
        if os.path.exists(fpath):
            try:
                with open(fpath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            except:
                with open(fpath, 'r', encoding='gbk', errors='ignore') as f:
                    content = f.read()
            
            content = clean_text(content)
            if content:
                all_text.append(content)
                print(f"[OK] 已读取: {os.path.basename(fpath)}")
    
    return '\n\n'.join(all_text)

def create_a4_doc(all_text):
    """创建A4文档"""
    doc = Document()
    
    # 设置A4页面
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # 封面
    title = doc.add_heading('外贸全链条培训资料', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('\n\n内部培训专用\n（A4打印版）')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    date_str = datetime.now().strftime('%Y年%m月%d日')
    date_para = doc.add_paragraph(f'\n\n生成日期：{date_str}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    print('[OK] 已添加封面')
    
    # 目录
    doc.add_heading('目录', 1)
    toc = [
        '一、外贸基础知识',
        '二、外贸人才与团队',
        '三、外贸获客渠道',
        '四、外贸谈判技巧',
        '五、跨境物流与通关',
        '六、外贸风险防范',
        '七、案例分析与实战',
        '八、常见问题解答',
        '九、工具与资源',
        '十、总结与行动计划',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    print('[OK] 已添加目录')
    
    # 正文 - 分成10章
    print('\n[INFO] 正在生成正文...')
    chunks = [all_text[i:i+800] for i in range(0, len(all_text), 800)]
    
    chapter_titles = [
        '外贸基础知识',
        '外贸人才与团队建设',
        '外贸获客渠道',
        '外贸谈判技巧',
        '跨境物流与通关',
        '外贸风险防范',
        '案例分析与实战',
        '常见问题解答',
        '工具与资源推荐',
        '总结与行动计划',
    ]
    
    total_chapters = min(10, len(chunks))
    chunks_per_chapter = max(1, len(chunks) // total_chapters)
    
    for i in range(total_chapters):
        start = i * chunks_per_chapter
        end = min((i+1) * chunks_per_chapter, len(chunks))
        chapter_text = '\n'.join(chunks[start:end])
        
        # 章节标题
        doc.add_heading(f'{i+1}. {chapter_titles[i]}', 1)
        
        # 章节内容
        paras = chapter_text.split('\n')
        for para_text in paras:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
        
        if i < total_chapters - 1:
            doc.add_page_break()
        
        print(f'[OK] 已生成章节 {i+1}: {chapter_titles[i]}')
    
    # 格式化
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = '微软雅黑'
        h_style.font.bold = True
        if i == 1:
            h_style.font.size = Pt(18)
        elif i == 2:
            h_style.font.size = Pt(14)
        else:
            h_style.font.size = Pt(12)
    
    print('\n[OK] 已设置文档格式')
    
    # 保存
    output = r'E:\郑州录音\外贸全链条培训资料_A4打印版.docx'
    doc.save(output)
    
    size_kb = os.path.getsize(output) / 1024
    print(f'\n{"="*60}')
    print(f'[SUCCESS] 文档已保存: {output}')
    print(f'[INFO] 文件大小: {size_kb:.1f} KB')
    print(f'[INFO] 总字符数: {len(all_text)}')
    print(f'{"="*60}')
    print('\n文档特点:')
    print('  • A4纸张 (21cm x 29.7cm)')
    print('  • 标准页边距 (上2.54cm, 下2.54cm, 左3.17cm, 右3.17cm)')
    print('  • 微软雅黑字体, 11pt, 1.5倍行距')
    print('  • 包含封面、目录、10个章节')
    print('  • 可直接打印!\n')

if __name__ == '__main__':
    print('='*60)
    print('开始生成A4打印版外贸培训资料')
    print('='*60)
    
    all_text = read_all_txt()
    
    if not all_text:
        print('[ERROR] 未找到txt文件!')
    else:
        print(f'\n[INFO] 总字符数: {len(all_text)}')
        create_a4_doc(all_text)
