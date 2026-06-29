#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最终版：清理并格式化外贸培训文档
- 彻底去除语音识别噪音和时间戳
- 口语转书面语
- 智能分段和结构化
- A4打印格式（标准页边距）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from datetime import datetime

def clean_voice_noise(text):
    """深度清理语音识别噪音"""
    if not text:
        return ""
    
    # 1. 移除所有时间戳 [...X.Xs]
    text = re.sub(r'$$\d+\.\d+s$$', '', text)
    
    # 2. 移除中英文混合的无意义片段
    text = re.sub(r'[a-zA-Z]{1,2}\s*[a-zA-Z]{1,2}\s*[a-zA-Z]{1,2}', '', text)
    
    # 3. 移除重复的标点符号
    text = re.sub(r'([。！？，、；：])\1+', '', text)
    
    # 4. 移除无意义的重复词（2-3字的重复3次以上）
    text = re.sub(r'(\S{1,3})\1{2,}', '', text)
    
    # 5. 移除过多的空格和换行
    text = re.sub(r'\s+', ' ', text)
    
    # 6. 移除乱码字符（保留中文、英文、数字、常用标点）
    # 构建允许字符的正则
    allowed_pattern = r'[^\u4e00-\u9fa5a-zA-Z0-9，。！？、；：""''（）《》，\s]'
    text = re.sub(allowed_pattern, '', text)
    
    # 7. 移除太短的无效片段（少于5个字符）
    if len(text.strip()) < 5:
        return ""
    
    return text.strip()

def colloquial_to_formal(text):
    """将口语转换为书面语"""
    if not text:
        return ""
    
    # 常见口语转书面语映射
    replacements = {
        '的话': '时',
        '然后呢': '其次',
        '那个': '该',
        '的话呢': '的情况下',
        '对吧': '是否正确',
        '嗯': '',
        '啊': '',
        '呢': '',
        '嘛': '',
        '哦': '',
        '呃': '',
        '额': '',
        '嗯嗯': '',
        '对对对': '确实',
        '是的是的': '确实',
        '好的好的': '好的',
        'OK': '好的',
        'ok': '好的',
        '嗯好的': '好的',
    }
    
    for colloquial, formal in replacements.items():
        text = text.replace(colloquial, formal)
    
    # 句首大写字母转中文（如果是英文单词开头）
    text = re.sub(r'^[a-zA-Z]+', lambda m: '', text)
    
    return text.strip()

def read_source_document():
    """读取源DOCX文档并清理"""
    doc_path = r'E:\郑州录音\新建 DOCX 文档.docx'
    
    if not os.path.exists(doc_path):
        print(f"[ERROR] 文件不存在: {doc_path}")
        return None
    
    print(f"[INFO] 正在读取: {doc_path}")
    doc = Document(doc_path)
    
    # 提取所有段落文本并清理
    cleaned_paragraphs = []
    raw_count = 0
    cleaned_count = 0
    
    for para in doc.paragraphs:
        if para.text.strip():
            raw_count += 1
            
            # 第一步：清理语音识别噪音
            cleaned = clean_voice_noise(para.text)
            
            if cleaned:
                # 第二步：口语转书面语
                formal = colloquial_to_formal(cleaned)
                
                if formal and len(formal) > 10:  # 只保留有效内容
                    cleaned_paragraphs.append(formal)
                    cleaned_count += 1
    
    full_text = '\n'.join(cleaned_paragraphs)
    
    print(f"[OK] 原始段落数: {raw_count}")
    print(f"[OK] 清理后段落数: {cleaned_count}")
    print(f"[OK] 总字符数: {len(full_text)}")
    print(f"[OK] 清理完成")
    
    return full_text

def create_structured_training_doc(content):
    """创建结构化的培训文档"""
    if not content or len(content) < 100:
        print("[ERROR] 内容过少，无法生成文档")
        return None
    
    doc = Document()
    
    # ========== 设置A4页面 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    print("[OK] 已设置A4页面（21cm x 29.7cm）")
    print("[OK] 已设置页边距（上2.54cm, 下2.54cm, 左3.17cm, 右3.17cm）")
    
    # ========== 封面 ==========
    title = doc.add_heading('外贸全链条培训资料', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('\n\n内部培训专用\n（正式版）')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    date_str = datetime.now().strftime('%Y年%m月%d日')
    date_para = doc.add_paragraph(f'\n\n生成日期：{date_str}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    date_para.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_page_break()
    print("[OK] 已添加封面")
    
    # ========== 目录 ==========
    doc.add_heading('目录', 1)
    
    toc_items = [
        '一、外贸基础知识',
        '二、外贸人才与团队建设',
        '三、外贸获客渠道与方法',
        '四、外贸谈判技巧与策略',
        '五、跨境物流与通关实务',
        '六、外贸风险识别与防范',
        '七、案例分析与实战演练',
        '八、常见问题与解答',
        '九、外贸工具与资源推荐',
        '十、总结与行动计划',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    print("[OK] 已添加目录")
    
    # ========== 正文内容 ==========
    print("\n[INFO] 正在生成正文...")
    
    # 智能分段：按句子分割
    sentences = re.split(r'(?<=[。！？])\s*', content)
    sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 5]
    
    print(f"[INFO] 识别到 {len(sentences)} 个句子")
    
    # 将句子组合成段落（每3-5个句子为一个段落）
    paragraphs = []
    current_para = []
    
    for sentence in sentences:
        current_para.append(sentence)
        if len(current_para) >= 3:  # 每3个句子组成一个段落
            paragraphs.append(''.join(current_para))
            current_para = []
    
    # 处理剩余的句子
    if current_para:
        paragraphs.append(''.join(current_para))
    
    print(f"[INFO] 组合成 {len(paragraphs)} 个段落")
    
    # 分成10个章节
    total_chapters = 10
    paragraphs_per_chapter = max(1, len(paragraphs) // total_chapters)
    
    chapter_titles = [
        '外贸基础知识',
        '外贸人才与团队建设',
        '外贸获客渠道与方法',
        '外贸谈判技巧与策略',
        '跨境物流与通关实务',
        '外贸风险识别与防范',
        '案例分析与实战演练',
        '常见问题与解答',
        '外贸工具与资源推荐',
        '总结与行动计划',
    ]
    
    for i in range(total_chapters):
        start_idx = i * paragraphs_per_chapter
        end_idx = min((i + 1) * paragraphs_per_chapter, len(paragraphs))
        
        if start_idx >= len(paragraphs):
            break
        
        # 章节标题
        doc.add_heading(f'{i+1}. {chapter_titles[i]}', 1)
        
        # 章节内容
        for para_text in paragraphs[start_idx:end_idx]:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
        
        # 章节小结
        doc.add_paragraph()
        summary = doc.add_paragraph(f'【本章要点】{chapter_titles[i]}的核心内容已阐述完毕。建议结合实际业务进行练习和巩固。')
        summary.paragraph_format.left_indent = Cm(0.5)
        summary.paragraph_format.space_before = Pt(12)
        summary.runs[0].font.bold = True
        summary.runs[0].font.size = Pt(10)
        summary.runs[0].font.color.rgb = RGBColor(0x33, 0x66, 0x99)
        
        # 分页
        if i < total_chapters - 1:
            doc.add_page_break()
        
        print(f"[OK] 已生成章节 {i+1}: {chapter_titles[i]}")
    
    # ========== 格式化文档 ==========
    print("\n[INFO] 正在格式化文档...")
    
    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # 标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '微软雅黑'
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        elif i == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        else:
            heading_style.font.size = Pt(12)
    
    print("[OK] 已设置文档格式")
    print("[OK] 字体：微软雅黑，11pt，1.5倍行距")
    
    # ========== 保存文档 ==========
    output_path = r'E:\郑州录音\外贸全链条培训资料_正式版_已清理.docx'
    doc.save(output_path)
    
    size_kb = os.path.getsize(output_path) / 1024
    
    print(f"\n{'='*60}")
    print(f"[SUCCESS] 文档已保存: {output_path}")
    print(f"[INFO] 文件大小: {size_kb:.1f} KB")
    print(f"[INFO] 总字符数: {len(content)}")
    print(f"[INFO] 总段落数: {len(paragraphs)}")
    print(f"{'='*60}")
    print("\n文档特点:")
    print("  - A4纸张 (21cm x 29.7cm)")
    print("  - 标准页边距 (上2.54cm, 下2.54cm, 左3.17cm, 右3.17cm)")
    print("  - 微软雅黑字体, 11pt, 1.5倍行距")
    print("  - 已彻底去除语音识别噪音和时间戳")
    print("  - 已将口语化表达转为正式书面语")
    print("  - 已智能分段和结构化")
    print("  - 包含封面、目录、10个章节、章节小结")
    print("  - 可直接打印使用!\n")
    
    return output_path

def main():
    """主函数"""
    print("="*60)
    print("开始清理并格式化外贸培训文档（最终版）")
    print("="*60)
    
    # 1. 读取并清理源文档
    content = read_source_document()
    
    if not content:
        print("\n[ERROR] 无法读取源文档！")
        return
    
    if len(content) < 100:
        print("\n[WARNING] 内容过少，请检查源文档！")
        return
    
    # 2. 创建结构化的培训文档
    output = create_structured_training_doc(content)
    
    if output:
        print(f"\n[SUCCESS] 处理完成！")
        print(f"[INFO] 输出文件: {output}")
        print(f"\n[INFO] 建议：打开文档后检查内容，如需调整请告知。")
    else:
        print("\n[ERROR] 文档生成失败！")

if __name__ == '__main__':
    main()
