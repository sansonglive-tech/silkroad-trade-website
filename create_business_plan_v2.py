# -*- coding: utf-8 -*-
"""
新丝路·企业主会 商业计划书 Word文档生成器 v2
基于截图内容整理，保持与图片一致的完整商业策划书格式
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_business_plan():
    doc = Document()
    
    # ==================== 页面设置 ====================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    
    # ==================== 封面 ====================
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("新丝路·企业主会")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("商业计划书")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 外贸企业数字化转型与出海一体化服务平台")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 底部信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("打造中原外贸新质生产力，赋能企业出海")
    run.font.size = Pt(14)
    run.font.italic = True
    
    # 分页
    doc.add_page_break()
    
    # ==================== 目录 ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目  录")
    run.font.size = Pt(24)
    run.font.bold = True
    
    doc.add_paragraph()
    
    toc_items = [
        "一、执行摘要",
        "二、市场机遇",
        "三、商业模式",
        "四、四层漏斗体系",
        "五、盈利测算",
        "六、实施路线图",
        "七、风险防范",
        "八、启动清单"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='Normal')
        p.paragraph_format.space_after = Pt(12)
        p.runs[0].font.size = Pt(14)
    
    doc.add_page_break()
    
    # ==================== 一、执行摘要 ====================
    add_heading(doc, "一、执行摘要", "003366")
    
    add_subheading(doc, "项目背景")
    add_paragraph_with_indent(doc, """新丝路·企业主会 是一个专为外贸中小企业打造的一站式出海服务平台，致力于通过数字化工具、系统化培训、高端人脉圈层和全链条服务，帮助企业快速打通外贸全链路，实现品牌出海。""")
    
    add_subheading(doc, "核心价值主张")
    
    # 价值主张表格
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["痛点", "解决方案"]
    values = [
        ["获客成本高、渠道单一", "四层漏斗体系，99元引流课精准获客"],
        ["团队不懂外贸、不会运营", "系统化培训+实操指导，手把手教会"],
        ["资源匮乏、抗风险能力弱", "高端人脉圈层+供应链整合赋能"],
        ["出海门槛高、不知从何入手", "交钥匙工程，从0到1全程陪跑"],
        ["传统模式效率低、增长慢", "数字化工具+AI提效，快速迭代增长"]
    ]
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for row_idx, (left, right) in enumerate(values):
        table.rows[row_idx + 1].cells[0].text = left
        table.rows[row_idx + 1].cells[1].text = right
        for cell in table.rows[row_idx + 1].cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    add_subheading(doc, "目标愿景")
    add_paragraph_with_indent(doc, """三年内成为中部地区最具影响力的外贸企业出海服务平台，帮助1000+企业实现品牌出海，年GMV突破10亿元。""")
    
    doc.add_page_break()
    
    # ==================== 二、市场机遇 ====================
    add_heading(doc, "二、市场机遇", "003366")
    
    add_subheading(doc, "宏观机遇")
    
    opportunities = [
        ("政策利好", "一带一路倡议持续深化，中原城市群崛起，郑州航空港、自贸区政策叠加优势明显"),
        ("市场需求", "传统外贸企业数字化转型需求迫切，跨境电商赛道持续高速增长"),
        ("技术赋能", "AI工具大幅降低外贸运营门槛，短视频/直播重构获客方式"),
        ("竞争格局", "市场上缺乏真正能提供全链条、一站式服务的综合平台")
    ]
    
    for title, desc in opportunities:
        p = doc.add_paragraph()
        run = p.add_run(f"▶ {title}：")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
        run = p.add_run(desc)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(8)
    
    add_subheading(doc, "目标客户画像")
    
    client_profiles = [
        "传统外贸企业老板，有产品但缺乏外贸运营能力",
        "制造业企业主，想拓展海外市场但不知如何入手",
        "跨境电商从业者，遇到瓶颈想突破但缺乏资源",
        "有产品、有供应链，想打造品牌出海的中小企业主"
    ]
    
    for i, profile in enumerate(client_profiles, 1):
        p = doc.add_paragraph(f"{i}. {profile}", style='List Number')
        p.runs[0].font.size = Pt(11)
    
    doc.add_page_break()
    
    # ==================== 三、商业模式 ====================
    add_heading(doc, "三、商业模式", "003366")
    
    add_paragraph_with_indent(doc, """新丝路·企业主会 采用"培训获客+社群沉淀+服务转化+生态增值"四维一体商业模式，构建从引流到转化到复购的完整商业闭环。""")
    
    add_subheading(doc, "收入来源")
    
    revenue_items = [
        ("前端引流产品", "99元引流课、299元入门课，低价获客，降低决策门槛"),
        ("中端转化产品", "付费社群（年费）、进阶培训课程，高性价比解决方案"),
        ("后端交付产品", "代运营服务、会销转化、定制化出海方案，高客单价值"),
        ("生态增值服务", "供应链整合、SaaS工具订阅、资源对接，持续复购")
    ]
    
    for title, desc in revenue_items:
        p = doc.add_paragraph()
        run = p.add_run(f"【{title}】")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
        run = p.add_run(desc)
        p.paragraph_format.space_after = Pt(10)
    
    doc.add_page_break()
    
    # ==================== 四、四层漏斗体系 ====================
    add_heading(doc, "四、四层漏斗体系", "003366")
    
    add_paragraph_with_indent(doc, """通过精心设计的四层漏斗体系，实现从海量流量到精准客户到高价值订单的层层转化。""")
    
    add_subheading(doc, "第一层：99元引流课")
    layers_1 = [
        "通过短视频/直播等渠道投放，吸引对外贸有兴趣的企业主",
        "以极低门槛切入，建立初步信任，展示专业能力",
        "目标：每月引流100-200人，转化率30%进入下一层"
    ]
    for item in layers_1:
        add_bullet_point(doc, item)
    
    add_subheading(doc, "第二层：进阶课程/付费社群")
    layers_2 = [
        "推出系列进阶课程（299-999元），深度讲解外贸全链路知识",
        "建立高端付费社群，年费3650-10000元，提供持续学习与人脉资源",
        "目标：转化率20-30%，筛选出真正有需求和付费能力的客户"
    ]
    for item in layers_2:
        add_bullet_point(doc, item)
    
    add_subheading(doc, "第三层：线下会销矩阵")
    layers_3 = [
        "举办线下培训会销活动（2-5万/人），面对面深度沟通",
        "邀请行业大咖、成功案例分享，打造高端圈层形象",
        "现场成单+后续跟进，目标客单价5-50万"
    ]
    for item in layers_3:
        add_bullet_point(doc, item)
    
    add_subheading(doc, "第四层：出海一体化服务平台")
    layers_4 = [
        "提供一站式出海服务：工商注册、海外合规、税务筹划、供应链整合等",
        "打造持续复购的生态体系，实现长期价值绑定",
        "目标：客户生命周期价值（LTV）10-50万"
    ]
    for item in layers_4:
        add_bullet_point(doc, item)
    
    doc.add_page_break()
    
    # ==================== 五、盈利测算 ====================
    add_heading(doc, "五、盈利测算", "003366")
    
    add_subheading(doc, "三年收入预测")
    
    # 盈利测算表格
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    
    headers = ["项目", "第一年", "第二年", "第三年", "累计"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    data = [
        ["付费用户数", "500人", "2000人", "5000人", "7500人"],
        ["客单价(平均)", "5000元", "8000元", "12000元", "-"],
        ["年收入", "250万", "1600万", "6000万", "7850万"],
        ["净利润率", "20%", "30%", "35%", "-"]
    ]
    
    for row_idx, row_data in enumerate(data):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = value
            for cell in table.rows[row_idx + 1].cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    add_subheading(doc, "成本结构")
    costs = [
        "人力成本：核心团队5-10人，年薪支出60-120万",
        "营销推广：短视频投放+内容制作，占营收15-20%",
        "场地设备：办公场地+线上工具，年支出10-20万",
        "课程研发：持续迭代课程体系，年投入20-30万",
        "其他费用：行政、法务、财务等，占营收5-8%"
    ]
    for cost in costs:
        add_bullet_point(doc, cost)
    
    doc.add_page_break()
    
    # ==================== 六、实施路线图 ====================
    add_heading(doc, "六、实施路线图", "003366")
    
    add_subheading(doc, "第一阶段：启动期（0-6个月）")
    phase1 = [
        "完成团队组建：运营、内容、销售、服务核心岗位到位",
        "开发核心产品：99元引流课+299元进阶课体系完成",
        "启动流量渠道：短视频账号矩阵搭建，持续输出内容",
        "建立社群模型：首批100人种子社群运营测试",
        "目标：完成MVP验证，月收入突破10万"
    ]
    for item in phase1:
        add_bullet_point(doc, item)
    
    add_subheading(doc, "第二阶段：成长期（6-18个月）")
    phase2 = [
        "扩大流量规模：多平台矩阵复制，付费投放放量",
        "完善产品矩阵：付费社群、进阶课程、会销活动全面推出",
        "打造成功案例：3-5个标杆客户，形成可复制方法论",
        "优化转化漏斗：各层级转化率持续优化提升",
        "目标：月收入突破100万，累计服务客户500+"
    ]
    for item in phase2:
        add_bullet_point(doc, item)
    
    add_subheading(doc, "第三阶段：规模化期（18-36个月）")
    phase3 = [
        "建立品牌影响力：成为区域外贸服务头部品牌",
        "拓展服务边界：从培训咨询延伸至代运营、出海服务",
        "技术工具加持：开发SaaS工具，提升服务效率",
        "生态联盟构建：链接供应链、服务商、渠道商",
        "目标：年营收突破5000万，净利润1500万+"
    ]
    for item in phase3:
        add_bullet_point(doc, item)
    
    doc.add_page_break()
    
    # ==================== 七、风险防范 ====================
    add_heading(doc, "七、风险防范", "003366")
    
    add_subheading(doc, "市场风险")
    add_bullet_point(doc, "风险：市场需求不及预期，付费转化率低")
    add_bullet_point(doc, "应对：持续打磨产品，建立成功案例库，灵活调整定价策略")
    
    add_subheading(doc, "竞争风险")
    add_bullet_point(doc, "风险：头部玩家入局，市场竞争加剧")
    add_bullet_point(doc, "应对：聚焦细分市场差异化，建立深度服务壁垒")
    
    add_subheading(doc, "运营风险")
    add_bullet_point(doc, "风险：团队扩张带来的管理挑战")
    add_bullet_point(doc, "应对：建立标准化流程，引入管理系统，梯队培养人才")
    
    add_subheading(doc, "政策风险")
    add_bullet_point(doc, "风险：跨境政策变化影响业务开展")
    add_bullet_point(doc, "应对：密切跟踪政策动向，建立合规体系，多元化业务布局")
    
    doc.add_page_break()
    
    # ==================== 八、启动清单 ====================
    add_heading(doc, "八、启动清单", "003366")
    
    add_subheading(doc, "立即行动（第一周）")
    actions1 = [
        "☐ 确定公司名称、完成工商注册",
        "☐ 注册微信公众号/视频号，完成基础搭建",
        "☐ 开发99元引流课，录制首批视频",
        "☐ 组建种子用户群（50-100人）"
    ]
    for action in actions1:
        p = doc.add_paragraph(action)
        p.runs[0].font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
    
    add_subheading(doc, "短期目标（第一个月）")
    actions2 = [
        "☐ 启动短视频/直播，每日更新内容",
        "☐ 举办首场直播公开课（99元引流）",
        "☐ 建立销售转化话术和流程",
        "☐ 完成首批付费用户转化（20-30人）"
    ]
    for action in actions2:
        p = doc.add_paragraph(action)
        p.runs[0].font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
    
    add_subheading(doc, "里程碑目标（三个月）")
    actions3 = [
        "☐ 月收入突破10万",
        "☐ 付费用户突破100人",
        "☐ 成功案例≥3个",
        "☐ 社群规模突破300人"
    ]
    for action in actions3:
        p = doc.add_paragraph(action)
        p.runs[0].font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
    
    # 封底
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("新丝路·企业主会")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 助力中国企业出海，共创全球贸易新篇章 ——")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # 保存文档
    output_path = r"E:\郑州录音\新丝路·企业主会_商业计划书_完整版.docx"
    doc.save(output_path)
    print(f"文档已生成: {output_path}")
    return output_path

def add_heading(doc, text, color_hex="003366"):
    """添加一级标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    return p

def add_subheading(doc, text):
    """添加二级标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_paragraph_with_indent(doc, text):
    """添加带缩进的段落"""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_bullet_point(doc, text):
    """添加项目符号"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

if __name__ == "__main__":
    create_business_plan()
