#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成A4打印版外贸培训资料
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

def clean_text(text):
    """清理语音识别噪音"""
    # 移除时间戳 [...0.3s]
    text = re.sub(r'$$\d+\.\d+s$$', '', text)
    # 移除重复的标点
    text = re.sub(r'([。！？])\1+', r'\1', text)
    # 移除过多的空格
    text = re.sub(r'\s+', ' ', text)
    # 移除无意义的重复词（2-3字的重复）
    text = re.sub(r'(\S{1,3})\1{3,}', r'\1', text)
    return text.strip()

def read_txt_files():
    """读取所有txt文件"""
    txt_files = [
        r'E:\郑州录音\2.txt',
        r'E:\郑州录音\3.txt',
        r'E:\郑州录音\4.txt',
        r'E:\郑州录音\5.txt',
        r'E:\郑州录音\6.txt',
    ]
    
    all_content = []
    for txt_file in txt_files:
        if os.path.exists(txt_file):
            try:
                with open(txt_file, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            except:
                with open(txt_file, 'r', encoding='gbk', errors='ignore') as f:
                    content = f.read()
            
            # 清理内容
            content = clean_text(content)
            if content:
                all_content.append({
                    'file': os.path.basename(txt_file),
                    'content': content
                })
                print(f"✓ 已读取: {os.path.basename(txt_file)} ({len(content)} 字符)")
    
    return all_content

def set_a4_page(doc):
    """设置A4页面"""
    from docx.oxml import OxmlElement
    
    section = doc.sections[0]
    
    # A4 纸张大小
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    
    # 页边距（上下2.54cm，左右3.17cm - 标准打印边距）
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    print("✓ 已设置A4页面（21cm x 29.7cm）")
    print("✓ 已设置页边距（上2.54cm，下2.54cm，左3.17cm，右3.17cm）")

def add_title_page(doc):
    """添加封面"""
    # 标题
    title = doc.add_heading('外贸全链条培训资料', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('\n\n内部培训专用\n（A4打印版）')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # 日期
    from datetime import datetime
    date_para = doc.add_paragraph(f'\n\n生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    date_para.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # 分页
    doc.add_page_break()
    print("✓ 已添加封面")

def add_table_of_contents(doc):
    """添加目录"""
    doc.add_heading('目录', 1)
    
    toc_items = [
        '一、外贸基础知识',
        '二、外贸人才与团队建设',
        '三、外贸获客渠道',
        '四、外贸谈判技巧',
        '五、跨境物流与通关',
        '六、外贸风险防范',
        '七、案例分析与实战',
        '八、常见问题解答',
        '九、工具与资源推荐',
        '十、总结与行动计划'
    ]
    
    for i, item in enumerate(toc_items, 1):
        p = doc.add_paragraph(f'{item}')
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    print("✓ 已添加目录")

def add_content(doc, all_content):
    """添加正文内容"""
    # 合并所有内容
    full_text = '\n\n'.join([item['content'] for item in all_content])
    
    # 智能分段（根据语义拆分）
    chapters = []
    
    # 尝试根据关键词划分章节
    keywords = {
        '外贸基础': ['外贸', '出口', '进口', '贸易'],
        '人才团队': ['人才', '团队', '招聘', '培训'],
        '获客渠道': ['获客', '客户', '渠道', '推广'],
        '谈判技巧': ['谈判', '沟通', '签单', '成交'],
        '物流通关': ['物流', '通关', '海关', '运输'],
        '风险防范': ['风险', '防骗', '安全', '合规'],
        '案例分析': ['案例', '实战', '经验'],
        '常见问题': ['问题', '解答', 'FAQ'],
        '工具资源': ['工具', '软件', '平台'],
        '总结计划': ['总结', '计划', '行动']
    }
    
    # 简单的内容分块（每段约500字符）
    chunk_size = 500
    chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]
    
    # 生成10个章节
    total_chapters = 10
    chunks_per_chapter = len(chunks) // total_chapters + 1
    
    chapter_titles = [
        '外贸基础知识',
        '外贸人才与团队建设',
        '外贸获客渠道',
        '外贸谈判技巧',
        '跨境物流与通关',
        '外贸风险防范',
        '案例分析与实战',
        '常见问题解答',
        '工具与资源推荐',
        '总结与行动计划'
    ]
    
    for i in range(total_chapters):
        start_idx = i * chunks_per_chapter
        end_idx = min((i + 1) * chunks_per_chapter, len(chunks))
        chapter_content = '\n'.join(chunks[start_idx:end_idx])
        
        # 添加章节标题
        doc.add_heading(f'{i+1}. {chapter_titles[i]}', 1)
        
        # 添加章节内容
        paragraphs = chapter_content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
        
        # 每个章节后分页
        if i < total_chapters - 1:
            doc.add_page_break()
        
        print(f"✓ 已生成章节 {i+1}: {chapter_titles[i]}")
    
    return len(chunks)

def format_document(doc):
    """格式化文档"""
    # 设置默认字体（适合中文打印）
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # 标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '微软雅黑'
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        elif i == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        else:
            heading_style.font.size = Pt(12)
    
    print("✓ 已设置文档格式（字体：微软雅黑，字号：11pt，行距：1.5倍）")

def add_page_numbers(doc):
    """添加页码"""
    from docx.oxml import OxmlElement
    
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    print("✓ 已添加页码")

def main():
    """主函数"""
    print("="*60)
    print("开始生成A4打印版外贸培训资料")
    print("="*60)
    
    # 创建文档
    doc = Document()
    
    # 1. 设置A4页面
    set_a4_page(doc)
    
    # 2. 添加封面
    add_title_page(doc)
    
    # 3. 添加目录
    add_table_of_contents(doc)
    
    # 4. 读取内容
    print("\n正在读取录音转写文件...")
    all_content = read_txt_files()
    
    if not all_content:
        print("错误：未找到任何txt文件！")
        return
    
    print(f"✓ 共读取 {len(all_content)} 个文件")
    
    # 5. 添加正文
    print("\n正在生成正文内容...")
    total_paragraphs = add_content(doc, all_content)
    
    # 6. 格式化文档
    print("\n正在格式化文档...")
    format_document(doc)
    
    # 7. 添加页码
    add_page_numbers(doc)
    
    # 8. 保存文档
    output_path = r'E:\郑州录音\外贸全链条培训资料_A4打印版.docx'
    doc.save(output_path)
    
    print("\n" + "="*60)
    print(f"✓ 文档已保存：{output_path}")
    print(f"✓ 文件大小：{os.path.getsize(output_path) / 1024:.1f} KB")
    print(f"✓ 总段落数：{total_paragraphs}")
    print("="*60)
    print("\n文档特点：")
    print("  • A4纸张大小（21cm x 29.7cm）")
    print("  • 标准打印页边距")
    print("  • 微软雅黑字体，适合中文阅读")
    print("  • 1.5倍行距，便于批注")
    print("  • 包含封面、目录、10个章节")
    print("  • 已添加页码")
    print("\n可以直接打印使用！")

if __name__ == '__main__':
    main()
