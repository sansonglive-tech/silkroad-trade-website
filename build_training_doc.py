# -*- coding: utf-8 -*-
"""
基于原始录音转写稿，重新整理成详尽的外贸培训学习资料
"""
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re

print("开始生成培训学习资料...")

# ========== 创建新文档 ==========
doc = docx.Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# ========== 封面 ==========
title = doc.add_heading('外贸全链条实战培训', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x6A)  # 深蓝色

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('内部学习资料 · 品牌运营官专项')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)  # 绿色

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('适用对象：外贸业务员 | 跨境电商运营 | 品牌运营官\n').font.size = Pt(10)
meta.add_run('编制日期：2026年6月').font.size = Pt(10)

doc.add_page_break()

# ========== 目录 ==========
doc.add_heading('目录', 1)
toc_items = [
    ('第一章', '外贸人才库搭建与管理', 'P5'),
    ('第二章', '客户付款方式与售后策略', 'P7'),
    ('第三章', '社媒账号安全与多账号管理', 'P9'),
    ('第四章', '定价策略与报价技巧', 'P11'),
    ('第五章', '链接优化与产品上架', 'P13'),
    ('第六章', '跨平台运营（国内/国外市场）', 'P15'),
    ('第七章', '生意方法论与长期主义', 'P17'),
    ('第八章', '实战案例与常见问题', 'P19'),
    ('第九章', '品牌运营官的一天（SOP）', 'P21'),
    ('第十章', '行动计划与考核标准', 'P23'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(f'{num}  {title}').bold = True
    p.add_run(f'  {page}')

doc.add_page_break()

# ========== 逐章生成内容 ==========

# ---------- 第一章 ----------
doc.add_heading('第一章  外贸人才库搭建与管理', 1)
doc.add_paragraph('【本章导读】人才是外贸业务的核心资产。本章讲解如何系统性搭建和管理外贸人才库，支撑业务持续增长。')

doc.add_heading('1.1 为什么需要外贸人才库', 2)
doc.add_paragraph('外贸业务具有周期长、环节多、专业门槛高的特点。一个完整的外贸人才库应包含：')
items = [
    '业务员：负责客户开发、谈判、跟单',
    '运营官：负责社媒引流、品牌塑造、内容输出',
    '供应链专员：负责工厂对接、质量管理、物流协调',
    '售后服务：负责客户维护、问题处理、复购引导',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.2 人才库搭建步骤', 2)
steps = [
    ('明确岗位需求', '根据业务阶段（起步期/成长期/成熟期）确定需要哪些岗位和人数。'),
    ('制定能力模型', '外贸业务员需要英语沟通+外贸流程知识；运营官需要社媒运营+内容创作能力。'),
    ('多渠道招聘', 'LinkedIn、外贸论坛、职业院校校企合作、内部推荐。'),
    ('培训体系搭建', '新人入职培训（产品知识+流程）+ 在职提升培训（谈判技巧+市场趋势）。'),
    ('考核与激励', '设定KPI（新客开发数、成交额、客户满意度），配套激励机制。'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    p.add_run(f'▶ {title}：').bold = True
    p.add_run(desc)

doc.add_heading('1.3 注意事项', 2)
notes = [
    '避免"人才闲置"：有明确的工作分配和考核，不让人才库变成"仓库"',
    '定期更新：业务变化后及时调整人才需求和能力模型',
    '知识沉淀：让老员工输出SOP和案例，变成新人培训材料',
]
for note in notes:
    doc.add_paragraph(note, style='List Number')

doc.add_paragraph('【本章小结】人才库不是静态名单，而是动态的能力管理系统。核心是：招得进、用得好、留得住。')
doc.add_page_break()

# ---------- 第二章 ----------
doc.add_heading('第二章  客户付款方式与售后策略', 1)
doc.add_paragraph('【本章导读】付款方式直接影响资金安全和客户体验，售后方式影响复购和口碑。两者需要系统设计方案。')

doc.add_heading('2.1 常见外贸付款方式对比', 2)
payment_table = doc.add_table(rows=5, cols=3)
payment_table.style = 'Light Grid Accent 1'
headers = ['付款方式', '适用场景', '风险提示']
for i, h in enumerate(headers):
    payment_table.rows[0].cells[i].text = h
    payment_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
data = [
    ('T/T 电汇', '大额定金+尾款，最常见', '需核实银行信息，防止诈骗'),
    ('L/C 信用证', '大额订单，银行担保', '单证要求严格，不符点会被拒付'),
    ('PayPal/信用卡', '小额订单，B2C业务', '手续费高，有chargeback风险'),
    ('西联/MoneyGram', '新兴市场，小额快速', '手续费高，有额度限制'),
]
for i, row_data in enumerate(data):
    for j, cell in enumerate(row_data):
        payment_table.rows[i+1].cells[j].text = cell

doc.add_heading('2.2 售后策略设计', 2)
doc.add_paragraph('售后不是"出问题才处理"，而是主动的客户关系管理：')
after_sales = [
    ('发货后24小时内', '发送发货通知+物流单号，让客户安心'),
    ('到货后3天', '主动询问收货情况，是否有质量问题'),
    ('使用1周后', '发送使用指南/温馨提示，提升体验'),
    ('使用1个月后', '询问反馈，引导复购或转介绍'),
    ('节假日', '发送祝福+新品推荐，保持联系'),
]
for time, action in after_sales:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{time}：').bold = True
    p.add_run(action)

doc.add_paragraph('【本章小结】付款方式选对，资金安全有保障；售后做主动，客户复购自然来。')
doc.add_page_break()

# ---------- 第三章 ----------
doc.add_heading('第三章  社媒账号安全与多账号管理', 1)
doc.add_paragraph('【本章导读】Facebook、WhatsApp、Google账号是外贸获客的重要工具，但平台严查多账号，一旦被封损失巨大。本章讲解安全操作方法。')

doc.add_heading('3.1 为什么账号会被封', 2)
reasons = [
    '同一设备/IP登录多个账号（平台通过设备指纹识别）',
    '账号注册信息重复（手机号、邮箱、身份证）',
    '操作行为异常（短时间内大量加好友、发消息）',
    '被用户举报（垃圾信息、骚扰）',
]
for i, r in enumerate(reasons):
    doc.add_paragraph(f'{i+1}. {r}', style='List Number')

doc.add_heading('3.2 多账号安全隔离方案', 2)
doc.add_paragraph('核心原则：每个账号 = 独立设备 + 独立IP + 独立身份信息')
solutions = [
    ('方案A：VPS/云服务器', '每个账号分配一台独立VPS，远程桌面操作，IP固定且独立'),
    ('方案B：指纹浏览器', '用AdsPower/Multilogin等工具，每个账号模拟独立设备指纹'),
    ('方案C：手机+流量卡', '实体手机+独立SIM卡，适合WhatsApp等需要手机验证的账号'),
]
for title, desc in solutions:
    p = doc.add_paragraph()
    p.add_run(f'▶ {title}：').bold = True
    p.add_run(desc)

doc.add_heading('3.3 日常操作规范', 2)
rules = [
    '新账号先"养号"7-14天：每天登录，浏览内容，少量互动，不要立即营销',
    '加好友要慢：每天不超过10个，分批添加，备注真实姓名',
    '内容发布：不要全部发广告，80%有价值内容+20%营销内容',
    '及时备份：账号密码、备用邮箱/手机号，记录在安全地方',
]
for rule in rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_paragraph('【本章小结】账号安全 = 技术隔离 + 行为规范。两者缺一不可。')
doc.add_page_break()

# ---------- 第四章 ----------
doc.add_heading('第四章  定价策略与报价技巧', 1)
doc.add_paragraph('【本章导读】定价直接影响利润和市场竞争力。报价技巧影响客户成交意愿。两者需要系统学习和实战训练。')

doc.add_heading('4.1 定价的三大方法', 2)
pricing = [
    ('成本加成定价', '公式：售价 = 成本 × (1 + 利润率)。适合标准化产品，缺点是没有考虑市场和竞争对手。'),
    ('竞争导向定价', '参考同行价格，制定略低或持平的价格。适合红海市场，利润较薄。'),
    ('价值导向定价', '根据客户感知价值定价，适合有品牌溢价或独特功能的产品。利润最高。'),
]
for title, desc in pricing:
    p = doc.add_paragraph()
    p.add_run(f'▶ {title}：').bold = True
    p.add_run(desc)

doc.add_heading('4.2 报价的5个技巧', 2)
tips = [
    ('阶梯报价', '数量越多，单价越低。引导客户下大单。例如：1-99个 $10/个；100-499个 $8.5/个；500+个 $7/个'),
    ('打包报价', '主产品+配件+售后打包，让客户感觉"超值"，同时提升客单价'),
    ('限时优惠', '"本周下单享5%折扣"，制造紧迫感，促进快速决策'),
    ('对比报价', '给出3个方案（基础/标准/高级），引导客户选中间档'),
    ('留谈判空间', '首次报价留10-15%的议价空间，让客户有"赢"的感觉'),
]
for i, (title, desc) in enumerate(tips):
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{title}：').bold = True
    p.add_run(desc)

doc.add_paragraph('【本章小结】定价是战略，报价是战术。定价定生死，报价定利润。')
doc.add_page_break()

# ---------- 第五章 ----------
doc.add_heading('第五章  链接优化与产品上架', 1)
doc.add_paragraph('【本章导读】无论是独立站、亚马逊、还是社媒带货，产品链接（Listing）质量直接决定转化率。本章讲解优化方法。')

doc.add_heading('5.1 产品链接的四大核心要素', 2)
elements = [
    ('标题', '包含核心关键词+核心卖点+适用场景。例如："2026新款便携养生壶 智能恒温 办公室家用 送礼佳品"'),
    ('主图', '白底高清图+场景图+细节图。主图决定点击率，附图决定转化率'),
    ('描述', '结构化呈现：产品参数 → 核心卖点 → 使用场景 → 售后保障。用短句、分点，便于手机阅读'),
    ('评价', '前10个评价决定后续转化。可以通过老客户、朋友、样品试用获取初始评价'),
]
for title, desc in elements:
    p = doc.add_paragraph()
    p.add_run(f'▶ {title}：').bold = True
    p.add_run(desc)

doc.add_heading('5.2 上架 checklist', 2)
checklist = [
    '□ 标题包含3个以上核心关键词',
    '□ 主图清晰，白底，无文字遮挡',
    '□ 附图至少5张（场景/细节/尺寸/包装/证书）',
    '□ 描述中有明确的"为什么买"和"为什么现在买"',
    '□ 价格有竞争力（对比前3名竞品）',
    '□ 库存数量真实，避免超卖',
    '□ 物流方案清晰（时效+费用）',
]
for item in checklist:
    doc.add_paragraph(item)

doc.add_paragraph('【本章小结】链接优化是持续过程，不是一次性工作。每周分析数据（点击率/转化率），持续优化。')
doc.add_page_break()

# ---------- 第六章 ----------
doc.add_heading('第六章  跨平台运营（国内/国外市场）', 1)
doc.add_paragraph('【本章导读】很多外贸人只关注国外市场，忽略了国内市场的机会。或者反过来。本章讲解如何统筹国内外两个市场。')

doc.add_heading('6.1 国内市场 vs 国外市场', 2)
compare_table = doc.add_table(rows=5, cols=3)
compare_table.style = 'Light Grid Accent 1'
headers = ['维度', '国内市场', '国外市场']
for i, h in enumerate(headers):
    compare_table.rows[0].cells[i].text = h
    compare_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
data = [
    ('客户特点', '决策快，重视关系', '决策慢，重视质量和认证'),
    ('沟通方式', '微信/电话，高效', '邮件/WhatsApp，正式'),
    ('付款方式', '支付宝/微信/银行转账', 'T/T/L/C/PayPal'),
    ('物流时效', '1-3天', '7-45天（看距离和方式）'),
]
for i, row_data in enumerate(data):
    for j, cell in enumerate(row_data):
        compare_table.rows[i+1].cells[j].text = cell

doc.add_heading('6.2 跨平台运营策略', 2)
strategies = [
    ('内容差异化', '国内发微信朋友圈/视频号，国外发TikTok/Facebook。内容风格要本地化。'),
    ('时间错峰', '国内白天工作时，正好是欧美晚上。可以请异地同事或外包团队覆盖时差。'),
    ('数据打通', '用CRM系统统一管理国内外客户，分析哪些产品在两个市场都受欢迎。'),
    ('供应链共享', '国内外订单可以共享供应链，提升采购规模，降低成本。'),
]
for title, desc in strategies:
    p = doc.add_paragraph()
    p.add_run(f'▶ {title}：').bold = True
    p.add_run(desc)

doc.add_paragraph('【本章小结】国内外市场不是"二选一"，而是"1+1>2"。核心是差异化运营，共享供应链。')
doc.add_page_break()

# ---------- 第七章 ----------
doc.add_heading('第七章  生意方法论与长期主义', 1)
doc.add_paragraph('【本章导读】做外贸不能只盯着眼前的订单，要有系统的方法论和长期的视角。本章分享经过验证的生意心法。')

doc.add_heading('7.1 做生意的三个原则', 2)
principles = [
    ('方法要对', '不是蛮干，而是用正确的方法做正确的事。例如：客户开发要用社媒+邮件组合，不能只靠等询盘。'),
    ('要对大家有利', '生意是共赢，不是零和游戏。让客户赚到钱，让供应商有利润，让自己有成长，这样的生意才能长久。'),
    ('要有利于事情本身', '做事情要问：这件事对客户有价值吗？对行业有推动吗？如果只是为了赚快钱，很难持续。'),
]
for title, desc in principles:
    p = doc.add_paragraph()
    p.add_run(f'▶ {title}：').bold = True
    p.add_run(desc)

doc.add_heading('7.2 长期主义的实践', 2)
long_term = [
    '不追求单笔订单利润最大化，追求客户终身价值（LTV）最大化',
    '持续投入品牌建设，哪怕短期看不到回报',
    '培养团队，让生意不依赖某个人',
    '关注行业趋势，提前布局下一个增长点',
]
for item in long_term:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('【本章小结】短期看技巧，长期看格局。外贸是一门生意，更是一门修行。')
doc.add_page_break()

# ---------- 第八章 ----------
doc.add_heading('第八章  实战案例与常见问题', 1)
doc.add_paragraph('【本章导读】通过真实案例学习，是最快的学习方式。本章分享典型成功案例和常见问题解决方案。')

doc.add_heading('8.1 案例：如何从0到月入10万', 2)
doc.add_paragraph('背景：某品牌运营官，之前没有外贸经验，用3个月时间实现月入10万。')
steps = [
    '第1个月：学习产品知识 + 搭建社媒账号（TikTok+Facebook），每天发1条内容',
    '第2个月：开始主动开发客户（通过TikTok私信+Facebook群组），累计开发200+潜客',
    '第3个月：成交12个客户，总销售额80万，利润约10万',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph('关键成功因素：', style='List Bullet')
for factor in ['选对了产品（有差异化，有利润空间）', '坚持发内容（没有因为前期没单就放弃）', '主动开发（不等询盘，主动找客户）']:
    doc.add_paragraph(factor, style='List Bullet')

doc.add_heading('8.2 常见问题 FAQ', 2)
faqs = [
    ('Q: 英语不好可以做外贸吗？', 'A: 可以。很多成功的外贸人英语也不好，用翻译工具+模板，关键是产品和服务。'),
    ('Q: 没有资金可以做外贸吗？', 'A: 可以从小额批发开始，或者做代发，不需要压库存。'),
    ('Q: 客户总是砍价怎么办？', 'A: 先了解客户砍价的原因（预算/比价/习惯），再针对性回应。不要直接拒绝，也不要无底线让步。'),
    ('Q: 账号被封了怎么办？', 'A: 立即申诉，同时准备备用账号。以后严格遵守操作规范，避免再次被封。'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    p.add_run(q + '\n').bold = True
    p.add_run(a)

doc.add_page_break()

# ---------- 第九章 ----------
doc.add_heading('第九章  品牌运营官的一天（SOP）', 1)
doc.add_paragraph('【本章导读】品牌运营官是外贸团队中的关键角色。本章给出标准工作日SOP，帮助新人快速上手。')

doc.add_heading('9.1 每日工作流程', 2)
schedule = [
    ('09:00-09:30', '查看数据：昨日曝光量、点击率、询盘数、成交量'),
    ('09:30-10:30', '内容创作：撰写今天要发布的社媒文案+挑选/制作图片'),
    ('10:30-12:00', '客户开发：通过TikTok/Facebook/WhatsApp主动联系潜客'),
    ('14:00-15:00', '客户跟进：回复询盘、报价、处理售后问题'),
    ('15:00-16:00', '内容发布：按预定时间发布社媒内容，互动回复评论'),
    ('16:00-17:00', '学习提升：看竞品、学新方法、优化话术'),
    ('17:00-18:00', '日报填写：记录今日工作成果、明日计划、遇到的问题'),
]
for time, task in schedule:
    p = doc.add_paragraph()
    p.add_run(f'{time}  ').bold = True
    p.add_run(task)

doc.add_heading('9.2 每周工作重点', 2)
weekly = [
    '周一：制定本周目标（新客开发数、内容发布数、成交额）',
    '周三：复盘前半周，调整策略',
    '周五：总结本周，规划下周，团队分享',
    '周末：竞品分析，内容素材收集',
]
for item in weekly:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('【本章小结】SOP不是束缚，而是保障。让每天的工作有重点、有节奏、有结果。')
doc.add_page_break()

# ---------- 第十章 ----------
doc.add_heading('第十章  行动计划与考核标准', 1)
doc.add_paragraph('【本章导读】学习要有输出，工作要有结果。本章给出具体的行动计划模板和考核标准，帮助落地执行。')

doc.add_heading('10.1 新人第一个月行动计划', 2)
plan_table = doc.add_table(rows=5, cols=3)
plan_table.style = 'Light Grid Accent 1'
headers = ['周次', '核心任务', '验收标准']
for i, h in enumerate(headers):
    plan_table.rows[0].cells[i].text = h
    plan_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
data = [
    ('第1周', '学习产品知识，搭建社媒账号', '能流利介绍产品，账号正常运作'),
    ('第2周', '开始发内容，学习客户开发话术', '发布7条内容，整理50个潜客名单'),
    ('第3周', '主动开发客户，回复询盘', '发送100+开发消息，回复所有询盘'),
    ('第4周', '跟进意向客户，促成成交', '至少1个成交客户，总结第一月经验'),
]
for i, row_data in enumerate(data):
    for j, cell in enumerate(row_data):
        plan_table.rows[i+1].cells[j].text = cell

doc.add_heading('10.2 考核标准（KPI）', 2)
doc.add_paragraph('品牌运营官的核心KPI：')
kpis = [
    '新客开发数：每月至少开发20个新客户（有有效沟通）',
    '内容发布数：每周至少发布5条社媒内容',
    '询盘回复率：100%（所有询盘24小时内回复）',
    '成交转化率：询盘→成交的转化率不低于5%',
    '客户满意度：售后满意度不低于90%',
]
for kpi in kpis:
    doc.add_paragraph(kpi, style='List Bullet')

doc.add_heading('10.3 激励方案', 2)
doc.add_paragraph('底薪 + 提成 + 奖金：')
incentives = [
    '底薪：保障基本生活',
    '提成：成交金额的2-5%（根据利润率调整）',
    '奖金：月度冠军奖、年度增长奖、团队贡献奖',
]
for item in incentives:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('【结语】外贸是一场马拉松，不是百米冲刺。系统学习、持续行动、长期坚持，一定能拿到结果。祝大家成功！')
doc.add_paragraph()
doc.add_paragraph('—— 品牌运营官培训项目组')
doc.add_paragraph('2026年6月')

# ========== 保存文档 ==========
output_path = r'E:\郑州录音\外贸全链条培训学习资料_最终版.docx'
doc.save(output_path)

print(f"SUCCESS: 文档已保存到 {output_path}")
print("文档包含：")
print("  - 封面 + 目录")
print("  - 10章完整内容（每章有导读、正文、小结）")
print("  - 表格、清单、FAQ、SOP、行动计划")
print("  - 总页数约25-30页")
