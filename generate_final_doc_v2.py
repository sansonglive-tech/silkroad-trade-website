# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

# 读取原始文档
print("Reading original DOCX...")
doc = docx.Document(r'E:\郑州录音\新建 DOCX 文档.docx')

# 提取所有段落文本
raw_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        raw_text.append(text)

print(f"Extracted {len(raw_text)} paragraphs")

# 清理录音停顿标记
def clean_text(text):
    # 移除 [...Xs] 这样的停顿标记
    text = re.sub(r'$$\.\.\.\d+\.?\d*s$$', '', text)
    # 移除其他可能的标记
    text = re.sub(r'$$\d+%$$', '', text)
    text = re.sub(r'$$\d+|$\d+$$', '', text)
    # 清理多余空格
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# 清理所有文本
cleaned_paragraphs = []
for text in raw_text:
    cleaned = clean_text(text)
    if cleaned:
        cleaned_paragraphs.append(cleaned)

print(f"Cleaned to {len(cleaned_paragraphs)} paragraphs")

# 创建新的学习资料文档
new_doc = docx.Document()

# 设置文档标题
title = new_doc.add_heading('外贸全链条培训学习资料', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加说明
new_doc.add_paragraph('本文档基于实战培训录音整理，涵盖外贸全链条运营、中医出海策略、品牌运营官技能等核心内容，适合外贸从业者、跨境电商运营人员系统学习。')
new_doc.add_paragraph()

# 简单的内容分块（每10段为一个章节）
chapter_size = 10
chapters = {}
for i, para_text in enumerate(cleaned_paragraphs):
    chapter_num = i // chapter_size + 1
    chapter_key = f'第{chapter_num}章：核心内容'
    if chapter_key not in chapters:
        chapters[chapter_key] = []
    chapters[chapter_key].append(para_text)

# 生成文档内容
for chapter_title, content_list in chapters.items():
    # 添加章节标题
    new_doc.add_heading(chapter_title, 1)
    
    # 添加章节内容
    for para_text in content_list:
        new_doc.add_paragraph(para_text)
    
    # 章节间添加空行
    new_doc.add_paragraph()

# 保存文档
output_path = r'E:\郑州录音\外贸全链条培训学习资料_最终版.docx'
new_doc.save(output_path)
print(f"SUCCESS! Document saved to: {output_path}")
print(f"Statistics:")
print(f"   - Total chapters: {len(chapters)}")
print(f"   - Total paragraphs: {sum(len(v) for v in chapters.values())}")
