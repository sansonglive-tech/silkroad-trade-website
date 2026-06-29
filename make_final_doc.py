# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 读取原始文档
doc = docx.Document(r'E:\郑州录音\新建 DOCX 文档.docx')

# 提取所有段落文本
raw_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        raw_text.append(text)

# 清理录音停顿标记
def clean_text(text):
    text = re.sub(r'$$\.\.\.\d+\.?\d*s$$', '', text)
    text = re.sub(r'$$\d+%$$', '', text)
    text = re.sub(r'$$\d+$$|$\d+$', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# 清理所有文本
cleaned_paragraphs = []
for text in raw_text:
    cleaned = clean_text(text)
    if cleaned:
        cleaned_paragraphs.append(cleaned)

# 创建新的学习资料文档
new_doc = docx.Document()

# 设置文档标题
title = new_doc.add_heading('外贸全链条培训学习资料', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加说明
new_doc.add_paragraph('本文档基于实战培训录音整理，涵盖外贸全链条运营、中医出海策略、品牌运营官技能等核心内容，适合外贸从业者、跨境电商运营人员系统学习。')
new_doc.add_paragraph()

# 手动划分章节（基于内容逻辑）
# 由于无法自动完美分类，这里按内容顺序分10章
chapter_size = len(cleaned_paragraphs) // 10 + 1

chapters = {}
for i, para_text in enumerate(cleaned_paragraphs):
    chapter_num = i // chapter_size + 1
    chapter_key = f'第{chapter_num}章：核心培训内容'
    if chapter_key not in chapters:
        chapters[chapter_key] = []
    chapters[chapter_key].append(para_text)

# 生成文档内容
for chapter_title, content_list in chapters.items():
    if content_list:
        new_doc.add_heading(chapter_title, 1)
        for para_text in content_list:
            new_doc.add_paragraph(para_text)
        new_doc.add_paragraph()

# 保存文档（先保存到工作区，避免权限问题）
import shutil
output_path = r'C:\Users\ASDCF\.qclaw\workspace\外贸全链条培训学习资料_最终版.docx'
new_doc.save(output_path)

# 尝试复制到目标目录
target_path = r'E:\郑州录音\外贸全链条培训学习资料_最终版.docx'
try:
    shutil.copy2(output_path, target_path)
    print('COPIED_TO_E_DRIVE')  # 简单输出，避免编码问题
except Exception as e:
    print('SAVED_ONLY_TO_WORKSPACE')  # 仅保存到工作区

exit(0)
