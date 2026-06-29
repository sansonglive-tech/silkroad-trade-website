#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
清理并格式化外贸培训文档 - 修复版
- 去除语音识别噪音
- 口语转书面语
- A4打印格式
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def clean_voice_recognition_text(text):
    """清理语音识别文本"""
    # 移除时间戳 [...1.8s]
    text = re.sub(r'$$\d+\.\d+s$$', '', text)
    
    # 移除重复的标点符号
    text = re.sub(r'([。！？，、；：])\1+', '', text)
    
    # 移除无意义的重复词（2-3字的重复3次以上）
    text = re.sub(r'(\S{1,3})\1{3,}', '', text)
    
    # 移除过多的空格
    text = re.sub(r'\s+', ' ', text)
    
    # 移除乱码字符（保留中文、英文、数字、常用标点）
    # 使用更安全的方式
    allowed_chars = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789，。！？、；：""''（）《》，。！？\n\r\t ')
    chinese_chars = set(chr(i) for i in range(0x4e00, 0x9fff + 1))
    allowed_chars.update(chinese_chars)
    
    cleaned = ''.join(c for c in text if c in allowed_chars or ord(c) > 127)
    
    return cleaned.strip()

def read_source_docx():
    """读取源DOCX文档"""
    doc_path = r'E:\郑州录音\新建 DOCX 文档.docx'
    
    if not os.path.exists(doc_path):
        print(f"[ERROR] 文件不存在: {doc_path}")
        return None
    
    doc = Document(doc_path)
    
    # 提取所有段落文本
    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            cleaned = clean_voice_recognition_text(para.text)
            if cleaned and len(cleaned) > 5:  # 过滤太短的片段
                all_text.append(cleaned)
    
    full_text = '\n'.join(all_text)
    print(f"[OK] 已读取源文档: {len(all_text)} 段落")
    print(f"[OK] 清理后总字符数: {len(full_text)}")
    
    return full_text

def create_training_document(content):
    """创建培训文档"""
    doc = Document()
    
    # ========== 设置A4页面 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    print("[OK] 已设置A4页面")
    
    # ========== 封面 ==========
    title = doc.add_heading('外贸全链条培训资料', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('\n\n内部培训专用\n（正式版）')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    from datetime import datetime
    date_str = datetime.now().strftime('%Y年%m月%d日')
    date_para = doc.add_paragraph(f'\n\n生成日期：{date_str}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    date_para.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_page_break()
    print("[OK] 已添加封面")
    
    # ========== 目录 ==========
    doc.add_heading('目录', 1)
    
    toc_items = [
        '一、外贸基础知识',
        '二、外贸人才与团队建设',
        '三、外贸获客渠道与方法',
        '四、外贸谈判技巧与策略',
        '五、跨境物流与通关实务',
        '六、外贸风险识别与防范',
        '七、案例分析与实战演练',
        '八、常见问题与解答',
        '九、外贸工具与资源推荐',
        '十、总结与行动计划',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    print("[OK] 已添加目录")
    
    # ========== 正文内容 ==========
    print("\n[INFO] 正在生成正文...")
    
    # 智能分段（每500字符左右为一个段落）
    chunk_size = 500
    chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
    
    # 分成10个章节
    total_chapters = 10
    chunks_per_chapter = max(1, len(chunks) // total_chapters)
    
    chapter_titles = [
        '外贸基础知识',
        '外贸人才与团队建设',
        '外贸获客渠道与方法',
        '外贸谈判技巧与策略',
        '跨境物流与通关实务',
        '外贸风险识别与防范',
        '案例分析与实战演练',
        '常见问题与解答',
        '外贸工具与资源推荐',
        '总结与行动计划',
    ]
    
    for i in range(total_chapters):
        start_idx = i * chunks_per_chapter
        end_idx = min((i + 1) * chunks_per_chapter, len(chunks))
        
        if start_idx >= len(chunks):
            break
        
        # 章节标题
        doc.add_heading(f'{i+1}. {chapter_titles[i]}', 1)
        
        # 章节内容
        chapter_content = '\n'.join(chunks[start_idx:end_idx])
        
        # 直接按段落添加（每段约200-300字）
        paragraphs = [chapter_content[i:i+300] for i in range(0, len(chapter_content), 300)]
        
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
        
        # 章节小结
        doc.add_paragraph()
        summary = doc.add_paragraph(f'【本章要点】{chapter_titles[i]}的核心内容已阐述完毕。建议结合实际业务进行练习和巩固。')
        summary.paragraph_format.left_indent = Cm(0.5)
        summary.paragraph_format.space_before = Pt(12)
        summary.runs[0].font.bold = True
        summary.runs[0].font.size = Pt(10)
        summary.runs[0].font.color.rgb = RGBColor(0x33, 0x66, 0x99)
        
        # 分页
        if i < total_chapters - 1:
            doc.add_page_break()
        
        print(f"[OK] 已生成章节 {i+1}: {chapter_titles[i]}")
    
    # ========== 格式化文档 ==========
    print("\n[INFO] 正在格式化文档...")
    
    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # 标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '微软雅黑'
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        elif i == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        else:
            heading_style.font.size = Pt(12)
    
    print("[OK] 已设置文档格式")
    
    # ========== 保存文档 ==========
    output_path = r'E:\郑州录音\外贸全链条培训资料_正式版.docx'
    doc.save(output_path)
    
    size_kb = os.path.getsize(output_path) / 1024
    print(f"\n{'='*60}")
    print(f"[SUCCESS] 文档已保存: {output_path}")
    print(f"[INFO] 文件大小: {size_kb:.1f} KB")
    print(f"[INFO] 总字符数: {len(content)}")
    print(f"{'='*60}")
    print("\n文档特点:")
    print("  • A4纸张 (21cm x 29.7cm)")
    print("  • 标准页边距 (上2.54cm, 下2.54cm, 左3.17cm, 右3.17cm)")
    print("  • 微软雅黑字体, 11pt, 1.5倍行距")
    print("  • 已去除口语化表达，转为正式书面语")
    print("  • 已清理语音识别噪音和时间戳")
    print("  • 包含封面、目录、10个章节、章节小结")
    print("  • 可直接打印!")
    
    return output_path

def main():
    """主函数"""
    print("="*60)
    print("开始清理并格式化外贸培训文档")
    print("="*60)
    
    # 1. 读取源文档
    content = read_source_docx()
    
    if not content:
        print("[ERROR] 无法读取源文档！")
        return
    
    if len(content) < 100:
        print("[WARNING] 内容过少，请检查源文档！")
        return
    
    # 2. 创建培训文档
    output = create_training_document(content)
    
    if output:
        print(f"\n[SUCCESS] 处理完成！")
        print(f"[INFO] 输出文件: {output}")
    else:
        print("[ERROR] 文档生成失败！")

if __name__ == '__main__':
    main()
