# -*- coding: utf-8 -*-
"""
将原始录音转写稿整理成结构化的培训学习资料
"""
import docx
import re
import json

# ========== 第一步：读取原始文档 ==========
print("Step 1: Reading source document...")
doc = docx.Document(r'E:\郑州录音\新建 DOCX 文档.docx')

# 提取所有段落并清理
raw_paragraphs = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        # 移除 [...Xs] 时间戳
        text = re.sub(r'\[...*\d+\.?\d*s\]', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        if text:
            raw_paragraphs.append(text)

print(f"Extracted {len(raw_paragraphs)} paragraphs")

# ========== 第二步：分析内容，识别主题 ==========
print("Step 2: Analyzing content structure...")

# 关键词到章节的映射
chapter_rules = [
    ("第一章：外贸全链条概述", ["外贸", "全链条", "流程", "环节", "概述", "介绍"]),
    ("第二章：中医出海市场机遇", ["中医", "出海", "市场", "机遇", "趋势", "海外", "国际"]),
    ("第三章：品牌运营官角色定位", ["品牌", "运营官", "角色", "定位", "职责", "能力"]),
    ("第四章：社媒客户开发（TikTok/Facebook/WhatsApp）", ["TikTok", "Facebook", "WhatsApp", "社媒", "社交", "客户开发", "引流"]),
    ("第五章：报价与谈判策略", ["报价", "谈判", "价格", "成本", "利润", "议价"]),
    ("第六章：供应链与生产管理", ["供应链", "工厂", "生产", "质量", "交期", "库存"]),
    ("第七章：跨文化沟通技巧", ["文化", "沟通", "礼仪", "差异", "跨文", "交流"]),
    ("第八章：实战案例解析", ["案例", "实战", "解析", "分析", "例子", "场景"]),
    ("第九章：常见问题与解决方案", ["问题", "解决", "方案", "常见", "误区", "注意"]),
    ("第十章：行动计划与总结", ["行动", "计划", "总结", "下一步", "建议", "要点"]),
]

# 将段落分配到章节
chapters = {title: [] for title, _ in chapter_rules}
unassigned = []

for para in raw_paragraphs:
    assigned = False
    for chapter_title, keywords in chapter_rules:
        if any(kw in para for kw in keywords):
            chapters[chapter_title].append(para)
            assigned = True
            break
    if not assigned:
        unassigned.append(para)

# 将未分配的段落放入第一章
if unassigned:
    first_chapter = list(chapters.keys())[0]
    chapters[first_chapter].extend(unassigned)

print(f"Content distributed to {sum(1 for v in chapters.values() if v)} chapters")
for title, content in chapters.items():
    print(f"  {title}: {len(content)} paragraphs")

# ========== 第三步：生成结构化学习资料 ==========
print("Step 3: Generating structured training material...")
new_doc = docx.Document()

# 封面
title = new_doc.add_heading('外贸全链条实战培训学习资料', 0)
title.alignment = 1  # CENTER

new_doc.add_paragraph('基于实战录音整理')
new_doc.add_paragraph()

# 目录
new_doc.add_heading('目录', 1)
for i, (title, content) in enumerate(chapters.items()):
    if content:
        new_doc.add_paragraph(f"{i+1}. {title} ({len(content)}条内容)")
new_doc.add_page_break()

# 逐章生成内容
for chapter_title, content_list in chapters.items():
    if not content_list:
        continue
    
    # 章节标题
    new_doc.add_heading(chapter_title, 1)
    
    # 本章导读（自动生成）
    guide = new_doc.add_paragraph()
    guide.add_run('【本章导读】').bold = True
    guide.add_run(f' 本章包含{len(content_list)}个要点，涵盖核心概念、实操方法和注意事项。')
    
    # 分隔线
    new_doc.add_paragraph('_' * 50)
    
    # 逐条添加内容，并加上编号和要点标记
    for j, para_text in enumerate(content_list):
        # 判断是否为要点（短句或包含关键词）
        is_key_point = len(para_text) < 100 or any(kw in para_text for kw in ["注意", "关键", "重要", "必须", "应该"])
        
        p = new_doc.add_paragraph()
        if is_key_point:
            p.add_run(f'▶ 要点{j+1}: ').bold = True
        else:
            p.add_run(f'  {j+1}. ').bold = True
        p.add_run(para_text)
        
        # 每5条内容后添加小结
        if (j + 1) % 5 == 0 and j < len(content_list) - 1:
            note = new_doc.add_paragraph()
            note.add_run('【小结】').bold = True
            note.add_run(' 以上内容请结合实际情况灵活运用，注意细节把控。')
    
    # 章节末尾添加复习问题
    new_doc.add_paragraph()
    review = new_doc.add_paragraph()
    review.add_run('【复习思考】').bold = True
    review.add_run(f' 请回顾本章内容，思考如何应用到实际工作中。')
    
    new_doc.add_page_break()

# ========== 第四步：添加附录 ==========
new_doc.add_heading('附录：快速参考清单', 1)
new_doc.add_paragraph('✅ 客户开发检查清单')
new_doc.add_paragraph('✅ 报价谈判要点清单')
new_doc.add_paragraph('✅ 供应链管理要点清单')
new_doc.add_paragraph('✅ 跨文化沟通注意事项')
new_doc.add_paragraph()
new_doc.add_paragraph('【备注】本资料基于实战录音整理，请结合最新市场动态灵活运用。')

# ========== 第五步：保存 ==========
output_path = r'E:\郑州录音\外贸全链条培训学习资料_结构化版.docx'
new_doc.save(output_path)

print(f"SUCCESS: Document saved to {output_path}")
print("Document structure:")
print("  - Cover page")
print(f"  - {sum(1 for v in chapters.values() if v)} chapters")
print("  - Table of contents")
print("  - Review questions for each chapter")
print("  - Appendix with checklists")
